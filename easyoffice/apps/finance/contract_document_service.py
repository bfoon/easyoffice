"""
apps/finance/contract_document_service.py
=========================================

Generates a Word (.docx) and PDF rendering of a Contract from its
structured fields, and provides helpers for the e-signature flow:

    * generate_contract_document(contract, actor) -> ContractDocument
    * register_uploaded_document(contract, actor) -> ContractDocument
    * build_signature_request(contract, document, signer_name, signer_email,
                              message, expires_in_days, actor)
        -> ContractSignatureRequest
    * stamp_signature_onto_pdf(signature_request, signature)
        -> writes signed PDF onto signature_request.signed_pdf
    * sha256_of_file(file_field) -> str

The DOCX is the editable master ("source of truth"). The PDF is the signing
copy and what the counterparty sees when they open the magic link.

Dependencies (already in your stack from letterhead/invoices):
    * python-docx
    * reportlab
    * Pillow (for signature image handling)
"""

from __future__ import annotations

import hashlib
import io
import os
import uuid
from datetime import datetime
from decimal import Decimal
from typing import Optional

from django.conf import settings
from django.core.files.base import ContentFile
from django.utils import timezone


# ────────────────────────────────────────────────────────────────────────────
# Errors
# ────────────────────────────────────────────────────────────────────────────

class ContractDocumentError(Exception):
    """Raised for predictable, user-presentable errors in this service."""


# ────────────────────────────────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────────────────────────────────

def sha256_of_file(file_field) -> str:
    """SHA-256 of a Django FileField's bytes, in hex."""
    if not file_field:
        return ''
    h = hashlib.sha256()
    file_field.open('rb')
    try:
        for chunk in file_field.chunks():
            h.update(chunk)
    finally:
        file_field.close()
    return h.hexdigest()


def _fmt_money(value, currency: str = 'GMD') -> str:
    try:
        v = Decimal(str(value or 0))
    except Exception:
        v = Decimal('0')
    return f'{currency} {v:,.2f}'


def _fmt_date(d) -> str:
    if not d:
        return '—'
    return d.strftime('%B %d, %Y')


def _build_snapshot(contract) -> dict:
    """Capture the contract's relevant fields at generation time."""
    return {
        'title':            contract.title,
        'reference':        contract.reference,
        'contract_type':    contract.get_contract_type_display(),
        'status':           contract.get_status_display(),
        'description':      contract.description,
        'start_date':       contract.start_date.isoformat() if contract.start_date else None,
        'end_date':         contract.end_date.isoformat() if contract.end_date else None,
        'renewal_date':     contract.renewal_date.isoformat() if contract.renewal_date else None,
        'standard_cost':    str(contract.standard_cost or 0),
        'currency':         contract.currency,
        'billing_cycle':    contract.get_billing_cycle_display(),
        'counterparty':     contract.counterparty_name,
        'counterparty_email': contract.counterparty_email,
        'vendor_company':   contract.vendor_company,
        'vendor_address':   contract.vendor_address,
        'vendor_phone':     contract.vendor_phone,
        'created_at':       timezone.now().isoformat(),
    }


def _branding():
    """Pull a few branding-y settings if you have them; fall back to safe defaults."""
    return {
        'org_name':   getattr(settings, 'OFFICE_NAME',   'Our Organisation'),
        'org_addr':   getattr(settings, 'OFFICE_ADDR',   ''),
        'org_phone':  getattr(settings, 'OFFICE_PHONE',  ''),
        'org_email':  getattr(settings, 'OFFICE_EMAIL',  ''),
        'org_color':  getattr(settings, 'OFFICE_COLOR',  '#0f172a'),
    }


# ────────────────────────────────────────────────────────────────────────────
# DOCX generation (python-docx)
# ────────────────────────────────────────────────────────────────────────────

def _build_docx_bytes(contract) -> bytes:
    """Build a polished Word document for `contract` and return its bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
    except ImportError as e:
        raise ContractDocumentError(
            'python-docx is not installed. Run: pip install python-docx'
        ) from e

    branding = _branding()
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Header band ────────────────────────────────────────────────────────
    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = head.add_run(branding['org_name'].upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    if branding['org_addr'] or branding['org_phone'] or branding['org_email']:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bits = [b for b in [branding['org_addr'], branding['org_phone'], branding['org_email']] if b]
        sub_run = sub.add_run(' · '.join(bits))
        sub_run.font.size = Pt(9)
        sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()  # spacer

    # ── Title ──────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run('CONTRACT AGREEMENT')
    tr.bold = True
    tr.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f'Reference: {contract.reference or "—"}')
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()  # spacer

    # ── Parties ────────────────────────────────────────────────────────────
    h = doc.add_paragraph()
    hr = h.add_run('THIS AGREEMENT IS MADE BETWEEN:')
    hr.bold = True
    hr.font.size = Pt(11)

    parties = doc.add_table(rows=2, cols=2)
    parties.autofit = True
    parties.columns[0].width = Cm(8.0)
    parties.columns[1].width = Cm(8.0)

    def _set_cell(cell, label, value, bold_value=True):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Clear default empty paragraph then add label + value
        cell.text = ''
        p1 = cell.paragraphs[0]
        r1 = p1.add_run(label)
        r1.font.size = Pt(8)
        r1.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        r1.bold = True
        p2 = cell.add_paragraph()
        r2 = p2.add_run(value or '—')
        r2.font.size = Pt(11)
        if bold_value:
            r2.bold = True

    _set_cell(parties.rows[0].cells[0], 'PARTY A (THE PROVIDER)', branding['org_name'])
    _set_cell(parties.rows[0].cells[1], 'PARTY B (THE COUNTERPARTY)', contract.counterparty_name)

    addr_a = '\n'.join([branding['org_addr'], branding['org_phone'], branding['org_email']]).strip('\n')
    addr_b = '\n'.join([
        contract.vendor_address,
        contract.vendor_phone,
        contract.counterparty_email,
    ]).strip('\n')

    _set_cell(parties.rows[1].cells[0], 'CONTACT', addr_a or '—', bold_value=False)
    _set_cell(parties.rows[1].cells[1], 'CONTACT', addr_b or '—', bold_value=False)

    doc.add_paragraph()  # spacer

    # ── Terms table ────────────────────────────────────────────────────────
    h = doc.add_paragraph()
    hr = h.add_run('1. KEY TERMS')
    hr.bold = True
    hr.font.size = Pt(12)

    terms_rows = [
        ('Contract type',       contract.get_contract_type_display()),
        ('Effective from',      _fmt_date(contract.start_date)),
        ('Expires on',          _fmt_date(contract.end_date)),
        ('Renewal checkpoint',  _fmt_date(contract.renewal_date)),
        ('Contract value',      _fmt_money(contract.standard_cost, contract.currency)),
        ('Billing cycle',       contract.get_billing_cycle_display()),
    ]

    tbl = doc.add_table(rows=len(terms_rows), cols=2)
    tbl.style = 'Light Grid Accent 1'
    for i, (k, v) in enumerate(terms_rows):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.text = ''
        c1.text = ''
        r = c0.paragraphs[0].add_run(k)
        r.bold = True
        r.font.size = Pt(10)
        r2 = c1.paragraphs[0].add_run(str(v))
        r2.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # ── Scope / description ────────────────────────────────────────────────
    h = doc.add_paragraph()
    hr = h.add_run('2. SCOPE OF AGREEMENT')
    hr.bold = True
    hr.font.size = Pt(12)

    desc = (contract.description or 'No additional scope details have been recorded.').strip()
    for chunk in desc.split('\n'):
        p = doc.add_paragraph(chunk)
        for r in p.runs:
            r.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # ── Standard clauses ───────────────────────────────────────────────────
    clauses = [
        ('3. PAYMENT TERMS',
         'Payment shall be made in accordance with the billing cycle and contract value '
         'set out above. Invoices issued under this agreement are payable within thirty '
         '(30) days of issuance unless agreed otherwise in writing.'),
        ('4. CONFIDENTIALITY',
         'Each party shall keep confidential all non-public information disclosed by the '
         'other in connection with this agreement, and shall use such information solely '
         'for the purpose of performing its obligations hereunder.'),
        ('5. TERM AND TERMINATION',
         'This agreement commences on the effective date and continues until the expiry '
         'date stated above, unless terminated earlier in accordance with its terms or '
         'extended by mutual written agreement of the parties.'),
        ('6. ELECTRONIC SIGNATURE',
         'The parties agree that this agreement may be executed by electronic signature '
         'and that an electronic signature shall have the same legal force and effect '
         'as a handwritten signature.'),
    ]
    for title, body in clauses:
        h = doc.add_paragraph()
        hr = h.add_run(title)
        hr.bold = True
        hr.font.size = Pt(12)
        p = doc.add_paragraph(body)
        for r in p.runs:
            r.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # ── Signature block ────────────────────────────────────────────────────
    h = doc.add_paragraph()
    hr = h.add_run('7. SIGNATURES')
    hr.bold = True
    hr.font.size = Pt(12)

    sig = doc.add_table(rows=4, cols=2)
    sig.autofit = True
    # Row 0: labels
    for i, label in enumerate(['PROVIDER (Party A)', 'COUNTERPARTY (Party B)']):
        c = sig.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(label)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Row 1: signature line placeholders
    for i in range(2):
        c = sig.rows[1].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run('\n\n_______________________________')
        r.font.size = Pt(10)

    # Row 2: name lines
    sig.rows[2].cells[0].text = ''
    sig.rows[2].cells[0].paragraphs[0].add_run(
        f'Name: {branding["org_name"]}'
    ).font.size = Pt(9)

    sig.rows[2].cells[1].text = ''
    sig.rows[2].cells[1].paragraphs[0].add_run(
        f'Name: {contract.counterparty_name}'
    ).font.size = Pt(9)

    # Row 3: date lines
    sig.rows[3].cells[0].text = ''
    sig.rows[3].cells[0].paragraphs[0].add_run('Date: ____________________').font.size = Pt(9)
    sig.rows[3].cells[1].text = ''
    sig.rows[3].cells[1].paragraphs[0].add_run('Date: ____________________').font.size = Pt(9)

    # ── Footer ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        f'Generated on {timezone.now().strftime("%B %d, %Y at %H:%M")} · '
        f'{contract.reference or "no reference"}'
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ────────────────────────────────────────────────────────────────────────────
# PDF generation (reportlab)
# ────────────────────────────────────────────────────────────────────────────

def _build_pdf_bytes(contract) -> bytes:
    """Build a signing-grade PDF for `contract`."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError as e:
        raise ContractDocumentError(
            'ReportLab is not installed. Run: pip install reportlab'
        ) from e

    branding = _branding()
    buf = io.BytesIO()
    pagesize = A4
    page_w, page_h = pagesize
    inner_w = page_w - 4 * cm

    doc = SimpleDocTemplate(
        buf,
        pagesize=pagesize,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=f'Contract — {contract.title}',
    )

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle('h1', parent=styles['Heading1'], alignment=TA_CENTER,
                        fontSize=20, leading=24, textColor=colors.HexColor('#0f172a'),
                        spaceAfter=6)
    sub = ParagraphStyle('sub', parent=styles['Normal'], alignment=TA_CENTER,
                         fontSize=10, leading=14, textColor=colors.HexColor('#64748b'),
                         spaceAfter=18)
    org = ParagraphStyle('org', parent=styles['Heading2'], alignment=TA_CENTER,
                         fontSize=13, leading=16, textColor=colors.HexColor('#0f172a'),
                         spaceAfter=4)
    org_sub = ParagraphStyle('org_sub', parent=styles['Normal'], alignment=TA_CENTER,
                             fontSize=8, leading=11, textColor=colors.HexColor('#64748b'),
                             spaceAfter=18)
    h2 = ParagraphStyle('h2', parent=styles['Heading2'], fontSize=11, leading=14,
                        textColor=colors.HexColor('#0f172a'), spaceBefore=14, spaceAfter=6)
    body = ParagraphStyle('body', parent=styles['Normal'], fontSize=10, leading=14,
                          textColor=colors.HexColor('#1f2937'), spaceAfter=4)
    label = ParagraphStyle('label', parent=styles['Normal'], fontSize=8, leading=11,
                           textColor=colors.HexColor('#64748b'))
    val = ParagraphStyle('val', parent=styles['Normal'], fontSize=11, leading=14,
                         textColor=colors.HexColor('#0f172a'))
    small = ParagraphStyle('small', parent=styles['Normal'], fontSize=8, leading=11,
                           alignment=TA_CENTER, textColor=colors.HexColor('#94a3b8'))

    story = []

    # Header
    story.append(Paragraph(branding['org_name'].upper(), org))
    bits = [b for b in [branding['org_addr'], branding['org_phone'], branding['org_email']] if b]
    if bits:
        story.append(Paragraph(' · '.join(bits), org_sub))

    story.append(Paragraph('CONTRACT AGREEMENT', h1))
    story.append(Paragraph(f'Reference: <b>{contract.reference or "—"}</b>', sub))

    # Parties block
    story.append(Paragraph('THIS AGREEMENT IS MADE BETWEEN:', h2))

    parties_data = [
        [
            Paragraph('PARTY A — PROVIDER', label),
            Paragraph('PARTY B — COUNTERPARTY', label),
        ],
        [
            Paragraph(f'<b>{branding["org_name"]}</b>', val),
            Paragraph(f'<b>{contract.counterparty_name or "—"}</b>', val),
        ],
        [
            Paragraph(
                ('<br/>'.join(filter(None, [
                    branding['org_addr'],
                    branding['org_phone'],
                    branding['org_email'],
                ]))) or '—',
                body,
            ),
            Paragraph(
                ('<br/>'.join(filter(None, [
                    contract.vendor_address,
                    contract.vendor_phone,
                    contract.counterparty_email,
                ]))) or '—',
                body,
            ),
        ],
    ]
    parties_tbl = Table(parties_data, colWidths=[inner_w / 2, inner_w / 2])
    parties_tbl.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('LINEBELOW', (0, 0), (-1, 0), 0.5, colors.HexColor('#e2e8f0')),
    ]))
    story.append(parties_tbl)

    # Key terms
    story.append(Paragraph('1. KEY TERMS', h2))
    terms = [
        ('Contract type',      contract.get_contract_type_display()),
        ('Effective from',     _fmt_date(contract.start_date)),
        ('Expires on',         _fmt_date(contract.end_date)),
        ('Renewal checkpoint', _fmt_date(contract.renewal_date)),
        ('Contract value',     _fmt_money(contract.standard_cost, contract.currency)),
        ('Billing cycle',      contract.get_billing_cycle_display()),
    ]
    terms_data = [
        [Paragraph(f'<b>{k}</b>', body), Paragraph(str(v), body)]
        for k, v in terms
    ]
    terms_tbl = Table(terms_data, colWidths=[5.5 * cm, inner_w - 5.5 * cm])
    terms_tbl.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#e2e8f0')),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f8fafc')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(terms_tbl)

    # Scope
    story.append(Paragraph('2. SCOPE OF AGREEMENT', h2))
    desc = (contract.description or 'No additional scope details have been recorded.').strip()
    for chunk in desc.split('\n'):
        if chunk.strip():
            story.append(Paragraph(chunk.replace('&', '&amp;'), body))

    # Clauses
    clauses = [
        ('3. PAYMENT TERMS',
         'Payment shall be made in accordance with the billing cycle and contract value set '
         'out above. Invoices issued under this agreement are payable within thirty (30) days '
         'of issuance unless agreed otherwise in writing.'),
        ('4. CONFIDENTIALITY',
         'Each party shall keep confidential all non-public information disclosed by the '
         'other in connection with this agreement, and shall use such information solely for '
         'the purpose of performing its obligations hereunder.'),
        ('5. TERM AND TERMINATION',
         'This agreement commences on the effective date and continues until the expiry date '
         'stated above, unless terminated earlier in accordance with its terms or extended by '
         'mutual written agreement of the parties.'),
        ('6. ELECTRONIC SIGNATURE',
         'The parties agree that this agreement may be executed by electronic signature and '
         'that an electronic signature shall have the same legal force and effect as a '
         'handwritten signature.'),
    ]
    for ttl, txt in clauses:
        story.append(Paragraph(ttl, h2))
        story.append(Paragraph(txt, body))

    # ── Signature page ────────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph('7. SIGNATURES', h2))
    story.append(Paragraph(
        'By signing below, each party agrees to be bound by the terms of this agreement.',
        body,
    ))
    story.append(Spacer(1, 1.2 * cm))

    # Two-column signature block. The right column is reserved for the
    # counterparty's signature stamp — its rectangle is at known coordinates
    # so we can stamp a PNG onto it later.
    sig_data = [
        [
            Paragraph('PROVIDER (Party A)', label),
            Paragraph('COUNTERPARTY (Party B)', label),
        ],
        [
            Paragraph('<br/><br/><br/>_______________________________', body),
            # We leave a visually-empty signature box; stamping happens
            # post-render. The "[signature here]" placeholder is just
            # advisory.
            Paragraph('<br/><br/><br/>_______________________________', body),
        ],
        [
            Paragraph(f'Name: <b>{branding["org_name"]}</b>', body),
            Paragraph(f'Name: <b>{contract.counterparty_name or "—"}</b>', body),
        ],
        [
            Paragraph('Date: ___________________', body),
            Paragraph('Date: ___________________', body),
        ],
    ]
    sig_tbl = Table(sig_data, colWidths=[inner_w / 2, inner_w / 2])
    sig_tbl.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(sig_tbl)

    story.append(Spacer(1, 1.5 * cm))
    story.append(Paragraph(
        f'Generated on {timezone.now().strftime("%B %d, %Y at %H:%M")} · '
        f'Reference {contract.reference or "—"}',
        small,
    ))

    doc.build(story)
    return buf.getvalue()


# ────────────────────────────────────────────────────────────────────────────
# Public API: document generation
# ────────────────────────────────────────────────────────────────────────────

def generate_contract_document(contract, *, actor=None):
    """
    Generate (or re-generate) a Word + PDF for `contract`.

    Returns the new ContractDocument row.
    """
    from apps.finance.models import ContractDocument

    # Build files
    docx_bytes = _build_docx_bytes(contract)
    pdf_bytes  = _build_pdf_bytes(contract)

    # Bump version: latest existing + 1
    last = ContractDocument.objects.filter(
        contract=contract,
        source=ContractDocument.Source.GENERATED,
    ).order_by('-version').first()
    next_version = (last.version + 1) if last else 1

    doc = ContractDocument(
        contract=contract,
        source=ContractDocument.Source.GENERATED,
        title=f'{contract.title} (v{next_version})',
        version=next_version,
        snapshot=_build_snapshot(contract),
        generated_by=actor,
    )

    # Save files. We use the contract reference as a filename hint when
    # available — easier for end users to recognise downloads.
    base_name = (contract.reference or f'contract-{contract.pk}').replace('/', '-')
    doc.docx_file.save(
        f'{base_name}-v{next_version}.docx',
        ContentFile(docx_bytes),
        save=False,
    )
    doc.pdf_file.save(
        f'{base_name}-v{next_version}.pdf',
        ContentFile(pdf_bytes),
        save=False,
    )
    doc.save()
    return doc


def register_uploaded_document(contract, *, actor=None):
    """
    Wrap the user-uploaded `contract.document` as a ContractDocument so the
    signing flow can target it the same way it targets generated documents.

    If `contract.document` is a PDF we use it directly as the PDF copy. If
    it's a DOCX we copy it as the DOCX copy and convert to PDF on the fly
    if possible; otherwise we keep the DOCX and skip PDF (signing a non-PDF
    is not supported by `stamp_signature_onto_pdf` — the UI must surface
    that limitation).
    """
    from apps.finance.models import ContractDocument

    if not contract.document:
        raise ContractDocumentError('This contract has no uploaded document.')

    last = ContractDocument.objects.filter(
        contract=contract,
        source=ContractDocument.Source.UPLOADED,
    ).order_by('-version').first()
    next_version = (last.version + 1) if last else 1

    name_lower = (contract.document.name or '').lower()
    is_pdf = name_lower.endswith('.pdf')
    is_docx = name_lower.endswith('.docx')

    doc = ContractDocument(
        contract=contract,
        source=ContractDocument.Source.UPLOADED,
        title=os.path.basename(contract.document.name),
        version=next_version,
        snapshot=_build_snapshot(contract),
        generated_by=actor,
    )

    # Read the raw bytes once.
    contract.document.open('rb')
    try:
        raw = contract.document.read()
    finally:
        contract.document.close()

    base_name = (contract.reference or f'contract-{contract.pk}').replace('/', '-')

    if is_pdf:
        doc.pdf_file.save(f'{base_name}-uploaded-v{next_version}.pdf',
                          ContentFile(raw), save=False)
    elif is_docx:
        doc.docx_file.save(f'{base_name}-uploaded-v{next_version}.docx',
                           ContentFile(raw), save=False)
        # Best-effort conversion via LibreOffice if available — silent fallback.
        try:
            pdf_bytes = _docx_to_pdf_via_libreoffice(raw)
            if pdf_bytes:
                doc.pdf_file.save(
                    f'{base_name}-uploaded-v{next_version}.pdf',
                    ContentFile(pdf_bytes),
                    save=False,
                )
        except Exception:
            # Conversion not available — that's fine, we keep the DOCX only.
            pass
    else:
        # Unknown type — store as PDF placeholder so download still works.
        doc.pdf_file.save(
            f'{base_name}-uploaded-v{next_version}{os.path.splitext(contract.document.name)[1]}',
            ContentFile(raw),
            save=False,
        )

    doc.save()
    return doc


def _docx_to_pdf_via_libreoffice(docx_bytes: bytes) -> Optional[bytes]:
    """Best-effort: convert DOCX → PDF using LibreOffice if it's installed."""
    import shutil
    import subprocess
    import tempfile

    soffice = shutil.which('libreoffice') or shutil.which('soffice')
    if not soffice:
        return None

    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, 'in.docx')
        with open(in_path, 'wb') as f:
            f.write(docx_bytes)
        try:
            subprocess.run(
                [soffice, '--headless', '--convert-to', 'pdf', '--outdir', td, in_path],
                check=True, timeout=60,
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            )
        except Exception:
            return None
        out_path = os.path.join(td, 'in.pdf')
        if os.path.exists(out_path):
            with open(out_path, 'rb') as f:
                return f.read()
    return None


# ────────────────────────────────────────────────────────────────────────────
# Signature requests
# ────────────────────────────────────────────────────────────────────────────

def build_signature_request(
    contract, document, *,
    signer_name: str,
    signer_email: str,
    message: str = '',
    expires_in_days: int = 30,
    actor=None,
):
    """
    Create a ContractSignatureRequest in DRAFT status.

    Caller is responsible for marking it as SENT (via mark_request_sent)
    after the email has been dispatched.
    """
    from datetime import timedelta
    from apps.finance.models import ContractSignatureRequest

    if not document.pdf_file:
        raise ContractDocumentError(
            'This document has no PDF available for signing. '
            'Generate the contract document first.'
        )

    sig_req = ContractSignatureRequest(
        contract=contract,
        document=document,
        signer_name=signer_name.strip(),
        signer_email=signer_email.strip(),
        message=message.strip(),
        status=ContractSignatureRequest.Status.DRAFT,
        expires_at=timezone.now() + timedelta(days=int(expires_in_days or 30)),
        created_by=actor,
    )

    # Take a tamper-evident hash of the PDF being signed.
    sig_req.document_hash = sha256_of_file(document.pdf_file)

    sig_req.save()
    return sig_req


def mark_request_sent(sig_req):
    from apps.finance.models import ContractSignatureRequest
    sig_req.status = ContractSignatureRequest.Status.SENT
    sig_req.sent_at = timezone.now()
    sig_req.save(update_fields=['status', 'sent_at', 'updated_at'])


def mark_request_viewed(sig_req):
    from apps.finance.models import ContractSignatureRequest
    now = timezone.now()
    fields = ['last_viewed_at', 'updated_at']
    if sig_req.first_viewed_at is None:
        sig_req.first_viewed_at = now
        fields.append('first_viewed_at')
    sig_req.last_viewed_at = now
    if sig_req.status == ContractSignatureRequest.Status.SENT:
        sig_req.status = ContractSignatureRequest.Status.VIEWED
        fields.append('status')
    sig_req.save(update_fields=fields)


# ────────────────────────────────────────────────────────────────────────────
# PDF stamping at signing time
# ────────────────────────────────────────────────────────────────────────────

def _typed_signature_to_png(text: str, size=(420, 130)) -> bytes:
    """
    Render typed-signature text as a transparent PNG using a cursive-ish font
    if available; falls back to a plain proportional font.
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError as e:
        raise ContractDocumentError(
            'Pillow is not installed. Run: pip install Pillow'
        ) from e

    img = Image.new('RGBA', size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)

    font = None
    for candidate in [
        '/usr/share/fonts/truetype/dejavu/DejaVuSerif-Italic.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSerif-Italic.ttf',
    ]:
        if os.path.exists(candidate):
            try:
                font = ImageFont.truetype(candidate, 64)
                break
            except Exception:
                continue
    if font is None:
        font = ImageFont.load_default()

    # Centre the text.
    bbox = draw.textbbox((0, 0), text, font=font)
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    x = (size[0] - tw) // 2
    y = (size[1] - th) // 2 - bbox[1]
    draw.text((x, y), text, fill=(15, 23, 42, 255), font=font)

    out = io.BytesIO()
    img.save(out, format='PNG')
    return out.getvalue()


def stamp_signature_onto_pdf(sig_req, signature) -> bytes:
    """
    Take the original PDF off `sig_req.document.pdf_file` and stamp the
    captured signature onto its last page (signature block area), then
    append a Completion Certificate page.

    Returns the bytes of the signed PDF, and ALSO writes them to
    `sig_req.signed_pdf`.
    """
    try:
        from pypdf import PdfReader, PdfWriter
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib import colors
    except ImportError as e:
        raise ContractDocumentError(
            'pypdf and reportlab are required. '
            'Run: pip install pypdf reportlab'
        ) from e

    src = sig_req.document.pdf_file
    if not src:
        raise ContractDocumentError('Source PDF is missing for this signature request.')

    src.open('rb')
    try:
        original_bytes = src.read()
    finally:
        src.close()

    reader = PdfReader(io.BytesIO(original_bytes))
    writer = PdfWriter()

    # Decide on the signature image bytes.
    if signature.method == signature.Method.TYPED:
        sig_png = _typed_signature_to_png(
            signature.typed_text or signature.signer_name_at_signing
        )
    else:
        # DRAWN or UPLOAD — use whatever is on the model.
        if not signature.signature_image:
            raise ContractDocumentError('Signature image is missing.')
        signature.signature_image.open('rb')
        try:
            sig_png = signature.signature_image.read()
        finally:
            signature.signature_image.close()

    # Build a stamp overlay sized to the LAST page (signature page).
    last_idx = len(reader.pages) - 1
    last_page = reader.pages[last_idx]
    pw = float(last_page.mediabox.width)
    ph = float(last_page.mediabox.height)

    # Stamp position: right column of the signature block.
    # The PDF was built with 2cm margins, so the right column starts at the
    # midpoint of the inner content width. We place the signature image
    # roughly where the right "________" line is.
    overlay_buf = io.BytesIO()
    c = canvas.Canvas(overlay_buf, pagesize=(pw, ph))

    # Compute placement (PDF origin is bottom-left).
    inner_left = 2 * cm
    inner_right = pw - 2 * cm
    inner_w = inner_right - inner_left
    col_left = inner_left + inner_w / 2 + 0.4 * cm
    col_right = inner_right
    sig_w = col_right - col_left - 0.5 * cm
    sig_h = sig_w * 0.32  # aspect ratio similar to a signature

    # The signature page from _build_pdf_bytes places the signature roughly
    # in the upper-middle of the page. We tune Y empirically so that the
    # signature image sits across the signature line. Adjust if your page
    # layout drifts.
    sig_x = col_left
    sig_y = ph - 9.5 * cm  # near the signature line region

    # Draw the signature image.
    from reportlab.lib.utils import ImageReader
    img_reader = ImageReader(io.BytesIO(sig_png))
    c.drawImage(img_reader, sig_x, sig_y, width=sig_w, height=sig_h,
                preserveAspectRatio=True, mask='auto')

    # Caption under the signature: name + signing timestamp.
    c.setFont('Helvetica', 8)
    c.setFillColor(colors.HexColor('#475569'))
    caption_y = sig_y - 0.55 * cm
    c.drawString(sig_x, caption_y, f'Signed by {signature.signer_name_at_signing}')
    c.drawString(sig_x, caption_y - 0.4 * cm,
                 timezone.now().strftime('%B %d, %Y at %H:%M UTC'))
    c.drawString(sig_x, caption_y - 0.8 * cm,
                 f'Method: {signature.get_method_display()}')
    c.save()
    overlay_buf.seek(0)

    overlay_reader = PdfReader(overlay_buf)
    overlay_page = overlay_reader.pages[0]

    # Copy each page from source — stamping only the last page.
    for i, page in enumerate(reader.pages):
        if i == last_idx:
            page.merge_page(overlay_page)
        writer.add_page(page)

    # Append a Completion Certificate page.
    cert_pdf = _build_completion_certificate(sig_req, signature)
    cert_reader = PdfReader(io.BytesIO(cert_pdf))
    for p in cert_reader.pages:
        writer.add_page(p)

    out_buf = io.BytesIO()
    writer.write(out_buf)
    signed_bytes = out_buf.getvalue()

    base_name = (sig_req.contract.reference or f'contract-{sig_req.contract_id}').replace('/', '-')
    sig_req.signed_pdf.save(
        f'{base_name}-signed-{sig_req.id.hex[:8]}.pdf',
        ContentFile(signed_bytes),
        save=False,
    )
    sig_req.status = sig_req.Status.SIGNED
    sig_req.signed_at = timezone.now()
    sig_req.save(update_fields=['signed_pdf', 'status', 'signed_at', 'updated_at'])

    return signed_bytes


def _build_completion_certificate(sig_req, signature) -> bytes:
    """A one-page audit certificate appended to every signed contract."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.enums import TA_CENTER

    branding = _branding()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title='Certificate of Completion',
    )
    styles = getSampleStyleSheet()
    title_st = ParagraphStyle('t', parent=styles['Heading1'], alignment=TA_CENTER,
                              fontSize=18, leading=22, textColor=colors.HexColor('#0f172a'))
    sub_st = ParagraphStyle('s', parent=styles['Normal'], alignment=TA_CENTER,
                            fontSize=10, leading=13, textColor=colors.HexColor('#64748b'),
                            spaceAfter=18)
    body_st = ParagraphStyle('b', parent=styles['Normal'], fontSize=10, leading=14,
                             textColor=colors.HexColor('#1f2937'))
    label_st = ParagraphStyle('l', parent=styles['Normal'], fontSize=8, leading=11,
                              textColor=colors.HexColor('#64748b'))
    val_st = ParagraphStyle('v', parent=styles['Normal'], fontSize=10, leading=14,
                            textColor=colors.HexColor('#0f172a'))

    inner_w = A4[0] - 4 * cm
    story = [
        Paragraph('CERTIFICATE OF COMPLETION', title_st),
        Paragraph(f'Issued by {branding["org_name"]}', sub_st),
    ]

    rows = [
        ('Contract',           sig_req.contract.title),
        ('Reference',          sig_req.contract.reference or '—'),
        ('Document version',   f'v{sig_req.document.version}'),
        ('Signer',             f'{signature.signer_name_at_signing} <{signature.signer_email_at_signing}>'),
        ('Signing method',     signature.get_method_display()),
        ('Signed at',          (sig_req.signed_at or timezone.now()).strftime('%B %d, %Y at %H:%M UTC')),
        ('Signer IP',          signature.signed_ip or '—'),
        ('User agent',         (signature.signed_user_agent or '—')[:120]),
        ('Document SHA-256',   (signature.document_hash_at_signing or sig_req.document_hash or '—')),
        ('Request ID',         str(sig_req.id)),
    ]
    table_data = [
        [Paragraph(k, label_st), Paragraph(str(v), val_st)]
        for k, v in rows
    ]
    tbl = Table(table_data, colWidths=[5.5 * cm, inner_w - 5.5 * cm])
    tbl.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#e2e8f0')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f8fafc')),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 0.7 * cm))
    story.append(Paragraph(
        'This certificate documents the electronic execution of the contract '
        'identified above. It is generated automatically and forms part of '
        'the contract record.',
        body_st,
    ))

    doc.build(story)
    return buf.getvalue()
