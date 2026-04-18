"""
apps/letterhead/docx_export.py

Pure-Python export — DOCX, PDF, PNG.
No Node.js required.

KEY DESIGN DECISIONS
────────────────────
• The exported files contain ONLY the letterhead header/footer — the body
  is left completely BLANK so the user can type their letter in Word or
  fill in their own content.  This matches the purpose of a letterhead.

• The header/footer rendering exactly mirrors the live HTML preview in
  builder.html for all 6 styles: classic, modern, bold, minimal, split,
  elegant.

• Three export formats:
    GET …/export/docx/  →  python-docx  (no Node.js)
    GET …/export/pdf/   →  reportlab    (no LibreOffice)
    GET …/export/png/   →  Pillow       (letterhead strip image)

Install (if not already):
    pip install python-docx reportlab Pillow

───────────────────────────────────────────────────────────────────────────
PDF HEADER FIX (April 2026)
───────────────────────────────────────────────────────────────────────────
The previous PDF export used a 50pt outer page margin and drew the top
accent bar *inside* that margin, which left white space around the
letterhead header and made the PDF look different from the live preview
in the builder.

This rewrite draws background shapes (accent bars, banners, panels)
edge-to-edge from x=0 to x=W, starting at y=H (the very top of the
page) — exactly matching the HTML preview where the header has no outer
margin.  Only the text/logo content respects an inner horizontal pad.
The footer is unchanged (it already worked correctly edge-to-edge).
"""

import io
import logging
import os

from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404
from django.views import View

from .models import LetterheadTemplate
from .views import _require_letterhead_permission

logger = logging.getLogger(__name__)


# ── Missing-dependency helper ─────────────────────────────────────────────
# Each exporter needs a specific third-party package. If a server doesn't
# have one installed, we raise this with a clear install hint so the view
# can return a friendly message instead of a raw ImportError.

class MissingDependencyError(RuntimeError):
    """Raised when a required export library isn't installed on the server."""
    pass


def _require(pkg_import: str, pip_name: str, feature: str):
    """
    Import a package at runtime and raise a friendly error if missing.

        doc_mod = _require('docx', 'python-docx', 'DOCX export')
    """
    import importlib
    try:
        return importlib.import_module(pkg_import)
    except ImportError as exc:
        raise MissingDependencyError(
            f"{feature} requires the '{pip_name}' Python package, which is "
            f"not installed on this server. Install it with: pip install {pip_name}"
        ) from exc


# ── Colour helpers ─────────────────────────────────────────────────────────

def _hex(h: str) -> str:
    """Strip # and return uppercase 6-char hex."""
    return h.lstrip("#").upper()

def _rgb(h: str) -> tuple:
    """'#1D9E75' → (29, 158, 117)"""
    h = _hex(h)
    return (int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def _rl_color(h: str):
    from reportlab.lib.colors import Color
    r,g,b = _rgb(h)
    return Color(r/255, g/255, b/255)


# ── Logo path helper ────────────────────────────────────────────────────────

def _logo_path(template: LetterheadTemplate):
    if not template.logo:
        return None
    try:
        return template.logo.path
    except (ValueError, NotImplementedError):
        import tempfile, uuid as _u
        ext = os.path.splitext(template.logo.name)[1] or ".png"
        p = os.path.join(tempfile.gettempdir(), f"lh_{_u.uuid4().hex}{ext}")
        try:
            with template.logo.open("rb") as s, open(p,"wb") as d:
                d.write(s.read())
            return p
        except Exception:
            return None


# ── HTML body parsing ──────────────────────────────────────────────────────
# The builder stores the user-edited body as HTML. For exports we need to
# convert that HTML back into a format the target engine understands.
#
# _strip_transient_html removes UI-only decorations (drag handles, toolbars,
# move-mode banner) that have no place in the exported output.
#
# _html_to_docx_paragraphs — walks the HTML, appending paragraphs and runs to
# a python-docx Document so headings, bold/italic/underline, and lists are
# preserved. Unsupported tags fall back to plain text.
#
# _html_to_pdf_flowables — produces a list of ReportLab flowables (Paragraph,
# Spacer, Table) that can be drawn into a PDF with a Frame.

def _strip_transient_html(html: str) -> str:
    """Remove builder UI chrome that shouldn't appear in exports."""
    if not html:
        return ""
    import re as _re
    # Remove move-mode info banner, element toolbars, drag handles
    html = _re.sub(r'<div[^>]*class=["\'][^"\']*\bmm-info\b[^"\']*["\'][^>]*>[\s\S]*?</div>',
                   '', html, flags=_re.I)
    html = _re.sub(r'<div[^>]*class=["\'][^"\']*\bel-toolbar\b[^"\']*["\'][^>]*>[\s\S]*?</div>',
                   '', html, flags=_re.I)
    html = _re.sub(r'<div[^>]*class=["\'][^"\']*\bmm-resize\b[^"\']*["\'][^>]*>[\s\S]*?</div>',
                   '', html, flags=_re.I)
    # Strip class/style attributes that are builder-specific
    html = _re.sub(r'\sdata-uid=["\'][^"\']*["\']', '', html, flags=_re.I)
    html = _re.sub(r'\sdata-draggable=["\'][^"\']*["\']', '', html, flags=_re.I)
    return html


# ── Bootstrap Icon → Unicode mapping ────────────────────────────────────────
# Bootstrap Icons uses a custom webfont. Since Word/PDF don't have that
# font, we substitute sensible Unicode characters for common icons. Any
# icon not in this map falls back to a bullet character ●.
#
# Coverage is pragmatic: the most common ~80 icons users insert in
# letterheads (checks, arrows, stars, common objects, social/contact
# symbols). Extend this map as needed.

_BI_TO_UNICODE = {
    # Checks & crosses
    "bi-check": "✓", "bi-check-lg": "✓", "bi-check2": "✓",
    "bi-check-circle": "✓", "bi-check-circle-fill": "✓",
    "bi-check-square": "☑", "bi-check-square-fill": "☑",
    "bi-x": "✗", "bi-x-lg": "✗", "bi-x-circle": "✗", "bi-x-circle-fill": "✗",
    # Stars & awards
    "bi-star": "☆", "bi-star-fill": "★", "bi-stars": "✨",
    "bi-trophy": "🏆", "bi-trophy-fill": "🏆", "bi-award": "🏅", "bi-award-fill": "🏅",
    # Arrows
    "bi-arrow-up": "↑", "bi-arrow-down": "↓", "bi-arrow-left": "←", "bi-arrow-right": "→",
    "bi-arrow-up-right": "↗", "bi-arrow-down-right": "↘",
    "bi-chevron-up": "‹", "bi-chevron-down": "›",
    "bi-arrow-right-circle": "➤", "bi-arrow-right-circle-fill": "➤",
    # Hearts & faces
    "bi-heart": "♡", "bi-heart-fill": "♥",
    "bi-emoji-smile": "🙂", "bi-emoji-smile-fill": "😀", "bi-emoji-heart-eyes": "😍",
    # Contact
    "bi-telephone": "☎", "bi-telephone-fill": "☎",
    "bi-envelope": "✉", "bi-envelope-fill": "✉", "bi-at": "@",
    "bi-globe": "🌐", "bi-globe2": "🌐",
    "bi-geo-alt": "📍", "bi-geo-alt-fill": "📍", "bi-pin-map": "📌", "bi-pin-map-fill": "📌",
    # Time & calendar
    "bi-clock": "⏰", "bi-clock-fill": "⏰", "bi-alarm": "⏰", "bi-alarm-fill": "⏰",
    "bi-calendar": "📅", "bi-calendar-fill": "📅",
    "bi-calendar-check": "📅", "bi-calendar-event": "📅",
    # Documents
    "bi-file": "📄", "bi-file-fill": "📄",
    "bi-file-earmark": "📄", "bi-file-earmark-fill": "📄",
    "bi-file-text": "📄", "bi-file-text-fill": "📄",
    "bi-files": "📑", "bi-folder": "📁", "bi-folder-fill": "📁",
    "bi-book": "📖", "bi-book-fill": "📖", "bi-journal-text": "📓",
    "bi-clipboard": "📋", "bi-clipboard-fill": "📋", "bi-clipboard-check": "📋",
    # People & orgs
    "bi-person": "👤", "bi-person-fill": "👤",
    "bi-people": "👥", "bi-people-fill": "👥",
    "bi-building": "🏢", "bi-building-fill": "🏢",
    "bi-bank": "🏛", "bi-bank2": "🏛",
    "bi-briefcase": "💼", "bi-briefcase-fill": "💼",
    # Money
    "bi-cash": "💵", "bi-cash-stack": "💵", "bi-currency-dollar": "$",
    "bi-credit-card": "💳", "bi-credit-card-fill": "💳",
    "bi-graph-up": "📈", "bi-graph-up-arrow": "📈", "bi-graph-down": "📉",
    # Tech & web
    "bi-wifi": "📶", "bi-shield": "🛡", "bi-shield-fill": "🛡",
    "bi-shield-check": "🛡", "bi-lock": "🔒", "bi-lock-fill": "🔒",
    "bi-unlock": "🔓", "bi-key": "🔑", "bi-key-fill": "🔑",
    # Misc common
    "bi-info-circle": "ℹ", "bi-info-circle-fill": "ℹ",
    "bi-exclamation-circle": "⚠", "bi-exclamation-triangle": "⚠",
    "bi-question-circle": "?", "bi-question-circle-fill": "?",
    "bi-lightbulb": "💡", "bi-lightbulb-fill": "💡",
    "bi-gear": "⚙", "bi-gear-fill": "⚙",
    "bi-flag": "⚑", "bi-flag-fill": "⚑",
    "bi-bell": "🔔", "bi-bell-fill": "🔔",
    "bi-bookmark": "🔖", "bi-bookmark-fill": "🔖",
    "bi-pencil": "✏", "bi-pencil-fill": "✏", "bi-pen": "🖊", "bi-pen-fill": "🖊",
    "bi-printer": "🖨", "bi-printer-fill": "🖨",
    "bi-camera": "📷", "bi-camera-fill": "📷",
    "bi-image": "🖼", "bi-image-fill": "🖼",
    "bi-music-note": "♪", "bi-music-note-beamed": "♫",
    "bi-search": "🔍",
    "bi-three-dots": "⋯", "bi-three-dots-vertical": "⋮",
    "bi-dot": "•", "bi-circle-fill": "●", "bi-circle": "○",
    "bi-square-fill": "■", "bi-square": "□",
    "bi-triangle-fill": "▲", "bi-triangle": "△",
    "bi-plus": "+", "bi-plus-lg": "+", "bi-plus-circle": "⊕", "bi-plus-circle-fill": "⊕",
    "bi-dash": "−", "bi-dash-lg": "−", "bi-dash-circle": "⊖",
    "bi-cart": "🛒", "bi-cart-fill": "🛒", "bi-bag": "🛍", "bi-bag-fill": "🛍",
    "bi-truck": "🚚", "bi-airplane": "✈", "bi-airplane-fill": "✈",
    "bi-house": "🏠", "bi-house-fill": "🏠",
    "bi-cup": "☕", "bi-cup-hot": "☕", "bi-cup-fill": "☕",
    # Social (use Unicode circled letters as placeholders — most
    # fonts won't have the brand glyphs; these are better than
    # missing-square boxes)
    "bi-facebook": "f", "bi-twitter": "𝕏", "bi-twitter-x": "𝕏",
    "bi-instagram": "📷", "bi-linkedin": "in", "bi-youtube": "▶",
    "bi-github": "⌨", "bi-whatsapp": "💬",
}
_FALLBACK_ICON_CHAR = "●"

def _bi_to_unicode(class_attr: str) -> str:
    """Return a Unicode character for the first bi-* class found, or a
    fallback bullet. Accepts the raw ``class="..."`` attribute value."""
    if not class_attr:
        return _FALLBACK_ICON_CHAR
    for cls in class_attr.split():
        if cls in _BI_TO_UNICODE:
            return _BI_TO_UNICODE[cls]
    return _FALLBACK_ICON_CHAR


def _replace_icons_with_unicode(html: str) -> str:
    """
    Replace every ``<i class="bi bi-..." ...></i>`` (Bootstrap Icon) with
    the matching Unicode character wrapped in a <span> that preserves its
    inline style (colour, font-size) so downstream parsers don't treat
    it as an italic element.
    """
    if not html:
        return ""
    import re as _re
    pattern = _re.compile(
        r'<i\b([^>]*\bclass\s*=\s*["\'][^"\']*\bbi\b[^"\']*["\'][^>]*)>[^<]*</i>',
        _re.I,
    )

    def _sub(m):
        attrs = m.group(1)
        # Pull class value
        cls_m = _re.search(r'class\s*=\s*["\']([^"\']*)["\']', attrs, _re.I)
        cls_val = cls_m.group(1) if cls_m else ""
        ch = _bi_to_unicode(cls_val)
        # Keep style if present (e.g. colour, size)
        style_m = _re.search(r'style\s*=\s*["\']([^"\']*)["\']', attrs, _re.I)
        style = style_m.group(1) if style_m else ""
        style_attr = f' style="{style}"' if style else ""
        return f'<span class="bi-icon"{style_attr}>{ch}</span>'

    return pattern.sub(_sub, html)


def _html_to_docx_paragraphs(html: str, doc, body_font: str = "Arial"):
    """
    Convert a chunk of HTML into python-docx paragraphs appended to ``doc``.
    Supports: <p>, <h1>-<h6>, <ul>/<ol>/<li>, <b>/<strong>, <i>/<em>,
    <u>, <br>, <hr>, <table>/<tr>/<td>/<th>, and Bootstrap Icons
    (<i class="bi bi-..."> → Unicode character).
    """
    if not html or not html.strip():
        return

    from html.parser import HTMLParser
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    html = _strip_transient_html(html)
    html = _replace_icons_with_unicode(html)

    # ── PASS 1: walk the DOM once to discover every <table> and count its
    # rows × columns so we can pre-create a correctly-sized docx table
    # (python-docx doesn't gracefully add columns after creation).
    class TableShapes(HTMLParser):
        def __init__(self):
            super().__init__(convert_charrefs=True)
            self.shapes = []      # list of (rows, max_cols) in document order
            self.stack  = []      # depth stack for nested tables (rare)
            self.cur_rows = 0
            self.cur_cols_in_row = 0
            self.cur_max_cols = 0
        def handle_starttag(self, tag, attrs):
            tag = tag.lower()
            if tag == 'table':
                self.stack.append({'rows':0, 'max_cols':0})
            elif tag == 'tr' and self.stack:
                self.stack[-1]['rows'] += 1
                self.stack[-1]['_cur_cols'] = 0
            elif tag in ('td','th') and self.stack:
                self.stack[-1]['_cur_cols'] = self.stack[-1].get('_cur_cols',0) + 1
                if self.stack[-1]['_cur_cols'] > self.stack[-1]['max_cols']:
                    self.stack[-1]['max_cols'] = self.stack[-1]['_cur_cols']
        def handle_endtag(self, tag):
            tag = tag.lower()
            if tag == 'table' and self.stack:
                info = self.stack.pop()
                self.shapes.append((max(1, info['rows']), max(1, info['max_cols'])))

    shape_walker = TableShapes()
    try:
        shape_walker.feed(html)
    except Exception:
        pass
    table_shapes = list(shape_walker.shapes)

    # ── PASS 2: emit paragraphs and tables
    class Emitter(HTMLParser):
        def __init__(self):
            super().__init__(convert_charrefs=True)
            self.style_stack = []
            self.current_p  = None
            self.list_stack = []         # stack of ('ul'|'ol', counter)
            self.default_size = 11
            self.skip_depth = 0
            # Nested-table support via stack
            self.table_stack = []        # list of dicts (see _start_table)
            self.table_shape_idx = 0     # which preview shape we're on

        def _active_cell(self):
            """If inside a table, return the current table cell; else None."""
            for tb in reversed(self.table_stack):
                if tb['cell'] is not None:
                    return tb['cell']
            return None

        def _start_para(self, size=None, align=None):
            if size is None:
                size = self.default_size
            cell = self._active_cell()
            if cell is not None:
                # First paragraph in a cell is the empty one docx creates.
                if cell.paragraphs and not cell.paragraphs[-1].text \
                        and not cell.paragraphs[-1].runs:
                    p = cell.paragraphs[-1]
                else:
                    p = cell.add_paragraph()
            else:
                p = doc.add_paragraph()
            if align:
                p.alignment = align
            p.paragraph_format.space_after = Pt(4)
            self.current_p = p
            self.current_size = size
            return p

        def _add_run(self, text):
            if not text:
                return
            if self.current_p is None:
                self._start_para()
            r = self.current_p.add_run(text)
            r.font.name = body_font
            r.font.size = Pt(self.current_size)
            if 'b' in self.style_stack: r.bold = True
            if 'i' in self.style_stack: r.italic = True
            if 'u' in self.style_stack: r.underline = True

        def _start_table(self):
            if self.table_shape_idx >= len(table_shapes):
                return None
            rows, cols = table_shapes[self.table_shape_idx]
            self.table_shape_idx += 1
            # Can't add tables inside a cell with doc.add_table — use the
            # cell's add_table method if we're nested.
            parent_cell = self._active_cell()
            if parent_cell is not None:
                try:
                    tbl = parent_cell.add_table(rows=rows, cols=cols)
                except TypeError:
                    tbl = parent_cell.add_table(rows=rows, cols=cols, width=None)
            else:
                tbl = doc.add_table(rows=rows, cols=cols)
            # Apply a light-grid style if it exists; fall back silently.
            try:
                tbl.style = 'Light Grid'
            except Exception:
                try: tbl.style = 'Table Grid'
                except Exception: pass
            self.table_stack.append({
                'tbl':    tbl,
                'rows':   rows,
                'cols':   cols,
                'r_idx': -1,      # index of current row (-1 = none yet)
                'c_idx':  0,
                'cell':   None,
                'header_row': False,
            })
            return tbl

        def _end_table(self):
            if self.table_stack:
                self.table_stack.pop()
            self.current_p = None

        def _start_row(self):
            if not self.table_stack:
                return
            tb = self.table_stack[-1]
            tb['r_idx'] += 1
            tb['c_idx'] = 0

        def _start_cell(self, is_header=False):
            if not self.table_stack:
                return
            tb = self.table_stack[-1]
            if tb['r_idx'] >= tb['rows'] or tb['c_idx'] >= tb['cols']:
                return   # guard against over-shoot
            cell = tb['tbl'].cell(tb['r_idx'], tb['c_idx'])
            tb['cell'] = cell
            tb['c_idx'] += 1
            tb['header_row'] = is_header
            # Reset current paragraph target; next content goes into this cell
            self.current_p = None
            if is_header:
                self.style_stack.append('b')

        def _end_cell(self, is_header=False):
            if not self.table_stack:
                return
            tb = self.table_stack[-1]
            tb['cell'] = None
            if is_header and 'b' in self.style_stack:
                self.style_stack.remove('b')
            self.current_p = None

        def handle_starttag(self, tag, attrs):
            tag = tag.lower()
            if self.skip_depth > 0:
                self.skip_depth += 1
                return
            if tag in ('script','style'):
                self.skip_depth = 1
                return
            if tag in ('b','strong'):
                # Skip when the <b> is inside an icon placeholder span
                self.style_stack.append('b')
            elif tag in ('i','em'):
                self.style_stack.append('i')
            elif tag == 'u':
                self.style_stack.append('u')
            elif tag == 'br':
                if self.current_p:
                    self.current_p.add_run().add_break()
            elif tag == 'p':
                self._start_para()
            elif tag in ('h1','h2','h3','h4','h5','h6'):
                sizes = {'h1':20,'h2':18,'h3':16,'h4':14,'h5':13,'h6':12}
                self.default_size = sizes[tag]
                self._start_para(size=sizes[tag])
                self.style_stack.append('b')
            elif tag == 'hr':
                self._start_para()
            elif tag in ('ul','ol'):
                self.list_stack.append([tag, 0])
            elif tag == 'li':
                self._start_para()
                if self.list_stack:
                    self.list_stack[-1][1] += 1
                    kind, idx = self.list_stack[-1]
                    bullet = '•  ' if kind == 'ul' else f'{idx}.  '
                    r = self.current_p.add_run(bullet)
                    r.font.name = body_font
                    r.font.size = Pt(self.current_size)
            elif tag == 'table':
                self._start_table()
            elif tag == 'tr':
                self._start_row()
            elif tag == 'td':
                self._start_cell(is_header=False)
            elif tag == 'th':
                self._start_cell(is_header=True)

        def handle_endtag(self, tag):
            tag = tag.lower()
            if self.skip_depth > 0:
                self.skip_depth -= 1
                return
            if tag in ('b','strong') and 'b' in self.style_stack:
                self.style_stack.remove('b')
            elif tag in ('i','em') and 'i' in self.style_stack:
                self.style_stack.remove('i')
            elif tag == 'u' and 'u' in self.style_stack:
                self.style_stack.remove('u')
            elif tag in ('h1','h2','h3','h4','h5','h6'):
                if 'b' in self.style_stack:
                    self.style_stack.remove('b')
                self.default_size = 11
            elif tag in ('ul','ol') and self.list_stack:
                self.list_stack.pop()
            elif tag == 'table':
                self._end_table()
            elif tag == 'td':
                self._end_cell(is_header=False)
            elif tag == 'th':
                self._end_cell(is_header=True)

        def handle_data(self, data):
            if self.skip_depth > 0:
                return
            if not data.strip():
                return
            text = ' '.join(data.split())
            self._add_run(text + ' ')

    try:
        Emitter().feed(html)
    except Exception as exc:
        logger.warning("HTML-to-DOCX fallback: %s", exc)
        import re as _re
        plain = _re.sub(r'<[^>]+>', ' ', html)
        plain = _re.sub(r'\s+', ' ', plain).strip()
        if plain:
            doc.add_paragraph(plain)


def _html_to_pdf_flowables(html: str, pri_color, font_h: str, font_b: str):
    """
    Convert HTML to ReportLab flowables. Emits Paragraph, Spacer, and
    Table objects. Supports the same subset as the DOCX variant.
    """
    if not html or not html.strip():
        return []
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import Paragraph, Spacer, Table as RLTable, TableStyle
    from reportlab.lib import colors as rl_colors

    html = _strip_transient_html(html)
    html = _replace_icons_with_unicode(html)

    styles = getSampleStyleSheet()
    normal   = ParagraphStyle('Body', parent=styles['Normal'],
                              fontName='Helvetica', fontSize=10, leading=14,
                              spaceAfter=3)
    heading1 = ParagraphStyle('H1', parent=normal, fontName='Helvetica-Bold',
                              fontSize=18, leading=22, spaceAfter=6)
    heading2 = ParagraphStyle('H2', parent=normal, fontName='Helvetica-Bold',
                              fontSize=14, leading=18, spaceAfter=5)

    from html.parser import HTMLParser
    flows = []

    class Walker(HTMLParser):
        def __init__(self):
            super().__init__(convert_charrefs=True)
            self.buffer = []
            self.current_style = normal
            self.list_stack = []          # ('ul'|'ol', counter)
            self.table_stack = []         # each: {'rows':[], 'cur_row':[], 'cur_cell':[]}
            self.skip = 0

        def _emit_paragraph(self, style=None):
            style = style or self.current_style
            txt = ''.join(self.buffer).strip()
            if txt:
                # Sanitize stray unmatched tags so Paragraph doesn't throw
                try:
                    flows.append(Paragraph(txt, style))
                except Exception:
                    import re as _re
                    safe = _re.sub(r'<(?!/?(b|i|u|br)\b)[^>]+>', '', txt)
                    try:
                        flows.append(Paragraph(safe, style))
                    except Exception:
                        flows.append(Paragraph(_re.sub(r'<[^>]+>', '', txt), normal))
                flows.append(Spacer(1, 3))
            self.buffer = []

        def _in_table(self):
            return bool(self.table_stack)

        def handle_starttag(self, tag, attrs):
            tag = tag.lower()
            if self.skip > 0:
                self.skip += 1; return
            if tag in ('script','style'):
                self.skip = 1; return
            if tag in ('b','strong'):
                self.buffer.append('<b>')
            elif tag in ('i','em'):
                self.buffer.append('<i>')
            elif tag == 'u':
                self.buffer.append('<u>')
            elif tag == 'br':
                self.buffer.append('<br/>')
            elif tag == 'p':
                self._emit_paragraph()
            elif tag == 'h1':
                self._emit_paragraph(); self.current_style = heading1
            elif tag == 'h2':
                self._emit_paragraph(); self.current_style = heading2
            elif tag in ('h3','h4','h5','h6'):
                self._emit_paragraph(); self.current_style = heading2
            elif tag == 'ul':
                self.list_stack.append(['ul', 0])
            elif tag == 'ol':
                self.list_stack.append(['ol', 0])
            elif tag == 'li':
                self._emit_paragraph()
                if self.list_stack:
                    self.list_stack[-1][1] += 1
                    kind, idx = self.list_stack[-1]
                    prefix = '•&nbsp;&nbsp;' if kind == 'ul' else f'{idx}.&nbsp;&nbsp;'
                    self.buffer.append(prefix)
            elif tag == 'hr':
                self._emit_paragraph()
                flows.append(Spacer(1, 6))
            elif tag == 'table':
                # Start a new table context. Emit anything pending first.
                self._emit_paragraph()
                self.table_stack.append({'rows': [], 'cur_row': None, 'cur_cell': None,
                                         'header_first_row': False, 'saw_header': False})
            elif tag == 'thead' and self._in_table():
                self.table_stack[-1]['header_first_row'] = True
            elif tag == 'tr' and self._in_table():
                self.table_stack[-1]['cur_row'] = []
            elif tag in ('td','th') and self._in_table():
                self.table_stack[-1]['cur_cell'] = []
                if tag == 'th':
                    self.table_stack[-1]['saw_header'] = True

        def handle_endtag(self, tag):
            tag = tag.lower()
            if self.skip > 0:
                self.skip -= 1; return
            if tag in ('b','strong'):
                self.buffer.append('</b>')
            elif tag in ('i','em'):
                self.buffer.append('</i>')
            elif tag == 'u':
                self.buffer.append('</u>')
            elif tag == 'p':
                self._emit_paragraph()
            elif tag == 'h1':
                self._emit_paragraph(heading1); self.current_style = normal
            elif tag == 'h2':
                self._emit_paragraph(heading2); self.current_style = normal
            elif tag in ('h3','h4','h5','h6'):
                self._emit_paragraph(heading2); self.current_style = normal
            elif tag == 'li':
                self._emit_paragraph()
            elif tag in ('ul','ol'):
                if self.list_stack:
                    self.list_stack.pop()
            elif tag in ('td','th') and self._in_table():
                tb = self.table_stack[-1]
                if tb['cur_cell'] is not None and tb['cur_row'] is not None:
                    cell_html = ''.join(tb['cur_cell']).strip() or '&nbsp;'
                    try:
                        p = Paragraph(cell_html, normal)
                    except Exception:
                        import re as _re
                        p = Paragraph(_re.sub(r'<[^>]+>', '', cell_html), normal)
                    tb['cur_row'].append(p)
                    tb['cur_cell'] = None
            elif tag == 'tr' and self._in_table():
                tb = self.table_stack[-1]
                if tb['cur_row']:
                    tb['rows'].append(tb['cur_row'])
                tb['cur_row'] = None
            elif tag == 'table' and self._in_table():
                tb = self.table_stack.pop()
                rows = tb['rows']
                if not rows:
                    return
                # Normalise column counts (pad short rows with empties)
                max_cols = max(len(r) for r in rows)
                for r in rows:
                    while len(r) < max_cols:
                        r.append(Paragraph('&nbsp;', normal))
                try:
                    rl_tbl = RLTable(rows, hAlign='LEFT', repeatRows=1 if tb['saw_header'] else 0)
                    styles_list = [
                        ('GRID', (0,0), (-1,-1), 0.5, rl_colors.grey),
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('LEFTPADDING', (0,0), (-1,-1), 4),
                        ('RIGHTPADDING', (0,0), (-1,-1), 4),
                        ('TOPPADDING', (0,0), (-1,-1), 3),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
                    ]
                    if tb['saw_header']:
                        styles_list += [
                            ('BACKGROUND', (0,0), (-1,0), rl_colors.whitesmoke),
                            ('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold'),
                        ]
                    rl_tbl.setStyle(TableStyle(styles_list))
                    flows.append(rl_tbl)
                    flows.append(Spacer(1, 6))
                except Exception as exc:
                    logger.warning("PDF table fallback: %s", exc)

        def handle_data(self, data):
            if self.skip > 0:
                return
            if not data.strip():
                return
            txt = ' '.join(data.split()) + ' '
            # Escape &, <, > for Paragraph safety (but our <b>/<i>/<u> are
            # raw strings in buffer — we escape only incoming text).
            esc = (txt.replace('&','&amp;')
                      .replace('<','&lt;').replace('>','&gt;'))
            tb = self.table_stack[-1] if self.table_stack else None
            if tb and tb.get('cur_cell') is not None:
                tb['cur_cell'].append(esc)
            else:
                self.buffer.append(esc)

    try:
        w = Walker(); w.feed(html); w._emit_paragraph()
    except Exception as exc:
        logger.warning("HTML-to-PDF fallback: %s", exc)
        import re as _re
        plain = _re.sub(r'<[^>]+>', ' ', html)
        plain = _re.sub(r'\s+', ' ', plain).strip()
        if plain:
            flows.append(Paragraph(plain, normal))

    return flows



# ═══════════════════════════════════════════════════════════════════════════
# DOCX  (python-docx)
# ═══════════════════════════════════════════════════════════════════════════

def _build_docx(t: LetterheadTemplate) -> bytes:
    _require('docx', 'python-docx', 'DOCX export')
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)

    # Header/footer edge-to-edge (flush with page top / bottom) — matches
    # the PDF export and the live preview. The *body* still has a normal
    # margin so the user has room to type a letter.
    sec.header_distance = Cm(0)
    sec.footer_distance = Cm(0)

    # top/bottom margin = where the body text starts/ends. Keep these
    # generous enough that body text doesn't collide with the
    # header/footer content (which sits in the 0..top_margin zone).
    sec.top_margin    = Cm(3.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.54)
    sec.right_margin  = Cm(2.54)

    pri = _hex(t.color_primary  or "#1D9E75")
    sec_col = _hex(t.color_secondary or "#2C2C2A")
    pri_rgb = RGBColor(*_rgb(t.color_primary  or "#1D9E75"))
    sec_rgb = RGBColor(*_rgb(t.color_secondary or "#2C2C2A"))
    fh = (t.font_header or "Georgia").strip("'")
    fb = (t.font_body   or "Arial"  ).strip("'")
    k  = t.template_key or "classic"

    hdr = sec.header
    hdr.is_linked_to_previous = False
    ftr = sec.footer
    ftr.is_linked_to_previous = False

    # ── helper: add a paragraph to a container ──────────────────────────
    def para(container, text="", bold=False, italic=False, size=11,
             color_rgb=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             font=None, sp_before=0, sp_after=4):
        p = container.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(sp_before)
        p.paragraph_format.space_after  = Pt(sp_after)
        if text:
            r = p.add_run(text)
            r.bold   = bold
            r.italic = italic
            r.font.name = font or fb
            r.font.size = Pt(size)
            if color_rgb:
                r.font.color.rgb = color_rgb
        return p

    # ── helper: horizontal rule on a paragraph (bottom border) ──────────
    def _hrule(p, color_hex, size=6, space=4, position="bottom"):
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        el   = OxmlElement(f"w:{position}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), str(space))
        el.set(qn("w:color"), color_hex)
        pBdr.append(el)
        pPr.append(pBdr)

    # ── helper: stretch a paragraph edge-to-edge ─────────────────────────
    # Uses negative indents to push the paragraph past the section's
    # left/right margins so its content (background shading, borders, and
    # text if you want) spans the full page width. The body text is
    # unaffected because it uses the normal section margins.
    #
    # Optional inner_pad_cm re-adds horizontal padding inside the stretched
    # paragraph so text doesn't sit flush against the paper edge, while
    # the background colour/border still extends to the edge.
    LEFT_M  = 2.54   # cm — matches sec.left_margin above
    RIGHT_M = 2.54
    def _edge_to_edge(p, inner_pad_cm=0.8):
        pf = p.paragraph_format
        pf.left_indent  = Cm(-LEFT_M  + inner_pad_cm)
        pf.right_indent = Cm(-RIGHT_M + inner_pad_cm)

    # ── helper: shade a paragraph background ────────────────────────────
    def _shade_para(p, fill_hex):
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill_hex)
        pPr.append(shd)

    # ── helper: shade a table cell ──────────────────────────────────────
    def _shade_cell(cell, fill_hex):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill_hex)
        tcPr.append(shd)

    # ── helper: set table cell borders to none ───────────────────────────
    def _no_borders(tbl):
        tbl_element = tbl._tbl
        tbl_pr = tbl_element.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl_element.insert(0, tbl_pr)
        bdr    = OxmlElement("w:tblBorders")
        for side in ("top","left","bottom","right","insideH","insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   "none")
            el.set(qn("w:sz"),    "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
            bdr.append(el)
        tbl_pr.append(bdr)

    # ── helper: stretch a table to full page width with negative indent
    # so the table's coloured cells / content reach the true page edges.
    # Full page = 21 cm; sets tblInd = -LEFT_M and tblW = 21 cm.
    def _full_width_table(tbl):
        tbl_element = tbl._tbl
        tbl_pr = tbl_element.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl_element.insert(0, tbl_pr)
        # Remove any existing tblInd/tblW
        for tag in ("w:tblInd","w:tblW"):
            existing = tbl_pr.find(qn(tag))
            if existing is not None:
                tbl_pr.remove(existing)
        # Negative indent to push table to left page edge (in twips)
        tbl_ind = OxmlElement("w:tblInd")
        tbl_ind.set(qn("w:w"),    str(-int(LEFT_M * 567)))
        tbl_ind.set(qn("w:type"), "dxa")
        tbl_pr.append(tbl_ind)
        # Width = 21 cm (A4 width) so it fills the whole page
        tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:w"),    str(int(21.0 * 567)))
        tbl_w.set(qn("w:type"), "dxa")
        tbl_pr.append(tbl_w)
        # Fixed layout so Word doesn't auto-resize cells
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_layout.set(qn("w:type"), "fixed")
        tbl_pr.append(tbl_layout)

    # ── helper: set per-cell internal margins (padding) in twips
    # default is ~108 twips (~0.19 cm) on left/right; we often want more.
    def _set_cell_margins(cell, top=100, left=200, bottom=100, right=200):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # Remove existing tcMar
        existing = tcPr.find(qn("w:tcMar"))
        if existing is not None:
            tcPr.remove(existing)
        mar = OxmlElement("w:tcMar")
        for side, val in (("top",top),("left",left),("bottom",bottom),("right",right)):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"),    str(val))
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        tcPr.append(mar)

    # ── helper: vertically centre cell content
    def _cell_vcenter(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        existing = tcPr.find(qn("w:vAlign"))
        if existing is not None:
            tcPr.remove(existing)
        v = OxmlElement("w:vAlign")
        v.set(qn("w:val"), "center")
        tcPr.append(v)

    # ── helper: put a coloured border on the LEFT side of a cell (accent stripe)
    def _cell_left_border(cell, color_hex, width_eighths_pt=24):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        existing = tcPr.find(qn("w:tcBorders"))
        if existing is not None:
            tcPr.remove(existing)
        bdr = OxmlElement("w:tcBorders")
        left = OxmlElement("w:left")
        left.set(qn("w:val"),   "single")
        left.set(qn("w:sz"),    str(width_eighths_pt))
        left.set(qn("w:space"), "0")
        left.set(qn("w:color"), color_hex)
        bdr.append(left)
        tcPr.append(bdr)

    # ── helper: try to insert logo image ────────────────────────────────
    def _add_logo(run, height_inches=0.5):
        lp = _logo_path(t)
        if lp:
            try:
                run.add_picture(lp, height=Inches(height_inches))
                return True
            except Exception:
                pass
        return False

    # ── address lines ────────────────────────────────────────────────────
    addr_lines = [l.strip() for l in (t.address or "").split("\n") if l.strip()]
    addr_flat  = "  ·  ".join(addr_lines)

    # ════════════════════════════════════════════════════════
    # CLASSIC  — top accent bar, name left, address right, footer band
    # ════════════════════════════════════════════════════════
    if k == "classic":
        # Top rule paragraph — edge-to-edge (inner_pad=0 so the rule spans
        # the entire page width, just like the PDF's c.rect(0, H-5, W, 5))
        p_top = para(hdr, sp_before=0, sp_after=4)
        _edge_to_edge(p_top, inner_pad_cm=0)
        _hrule(p_top, pri, size=18, space=2, position="top")

        # Logo + name
        p_name = para(hdr, sp_before=4, sp_after=2)
        r_logo = p_name.add_run()
        has_logo = _add_logo(r_logo, 0.5)
        if has_logo:
            p_name.add_run("  ")
        r_name = p_name.add_run(t.company_name or "")
        r_name.bold = True; r_name.font.name = fh; r_name.font.size = Pt(20)
        r_name.font.color.rgb = sec_rgb

        if t.tagline:
            p_tag = para(hdr, t.tagline, italic=True, size=11, color_rgb=pri_rgb, sp_after=2)

        for line in addr_lines:
            para(hdr, line, size=9, color_rgb=RGBColor(0x88,0x88,0x88),
                 align=WD_ALIGN_PARAGRAPH.RIGHT, sp_after=0)
        if t.contact_info:
            para(hdr, t.contact_info, size=9, color_rgb=RGBColor(0xaa,0xaa,0xaa),
                 align=WD_ALIGN_PARAGRAPH.RIGHT, sp_after=4)

        # Divider
        p_div = para(hdr, sp_before=2, sp_after=0)
        _hrule(p_div, "DDDDDD", size=4, space=2)

        # Footer band — edge-to-edge coloured band with text inside
        p_ftr = para(ftr, f"{t.contact_info or ''}  ·  {addr_flat}",
                     size=9, color_rgb=RGBColor(255,255,255),
                     align=WD_ALIGN_PARAGRAPH.CENTER, sp_before=4, sp_after=4)
        _edge_to_edge(p_ftr, inner_pad_cm=0)
        _shade_para(p_ftr, pri)
        _hrule(p_ftr, pri, size=6, space=4, position="top")

    # ════════════════════════════════════════════════════════
    # MODERN  — left accent border, name, contact right, thin rule
    # ════════════════════════════════════════════════════════
    # ════════════════════════════════════════════════════════
    # MODERN  — flush-left accent stripe, name + tagline, contact right
    # ════════════════════════════════════════════════════════
    elif k == "modern":
        from docx.shared import Cm as _Cm
        # 3-column table, full page width:
        #   col 0 = narrow shaded stripe (the accent)
        #   col 1 = content (name, tagline, address)
        #   col 2 = tiny right-side gutter
        tbl = hdr.add_table(rows=1, cols=3, width=_Cm(21))
        _no_borders(tbl)
        _full_width_table(tbl)
        tbl.columns[0].width = _Cm(0.35)
        tbl.columns[1].width = _Cm(20.35)
        tbl.columns[2].width = _Cm(0.30)
        c_stripe = tbl.cell(0,0)
        c_main   = tbl.cell(0,1)
        c_right  = tbl.cell(0,2)
        # Shaded narrow cell = the accent stripe (guaranteed-visible)
        _shade_cell(c_stripe, pri)
        _set_cell_margins(c_stripe, top=0, left=0, bottom=0, right=0)
        _set_cell_margins(c_main,   top=240, left=320, bottom=200, right=0)
        _set_cell_margins(c_right,  top=0, left=0, bottom=0, right=0)

        # Main cell: name line (with optional logo)
        p_name = c_main.paragraphs[0]
        p_name.paragraph_format.space_after = Pt(2)
        r_logo = p_name.add_run()
        has_logo = _add_logo(r_logo, 0.45)
        if has_logo:
            p_name.add_run("  ")
        r_name = p_name.add_run(t.company_name or "")
        r_name.bold = True; r_name.font.name = fh; r_name.font.size = Pt(18)
        r_name.font.color.rgb = sec_rgb

        if t.tagline:
            p_tag = c_main.add_paragraph()
            p_tag.paragraph_format.space_after = Pt(2)
            r_t = p_tag.add_run(t.tagline)
            r_t.italic = True; r_t.font.size = Pt(10)
            r_t.font.color.rgb = RGBColor(0xaa,0xaa,0xaa)

        # Contact + address, right-aligned
        if t.contact_info:
            p_c = c_main.add_paragraph()
            p_c.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_c.paragraph_format.space_after = Pt(0)
            # Negative right-indent on this paragraph to push up to right edge
            p_c.paragraph_format.right_indent = Cm(-0.5)
            r_c = p_c.add_run(t.contact_info)
            r_c.font.size = Pt(9)
            r_c.font.color.rgb = RGBColor(0x88,0x88,0x88)
        for line in addr_lines:
            p_l = c_main.add_paragraph()
            p_l.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_l.paragraph_format.space_after = Pt(0)
            p_l.paragraph_format.right_indent = Cm(-0.5)
            r_l = p_l.add_run(line)
            r_l.font.size = Pt(9)
            r_l.font.color.rgb = RGBColor(0xbb,0xbb,0xbb)

        # Thin divider under the whole header (edge-to-edge)
        p_div = para(hdr, sp_before=4, sp_after=0)
        _edge_to_edge(p_div, inner_pad_cm=0)
        _hrule(p_div, "E0E0E0", size=4, space=2)

        # Footer rule spans the full width
        p_ftr = para(ftr, f"{t.contact_info or ''}  ·  {addr_flat}",
                     size=9, color_rgb=pri_rgb,
                     align=WD_ALIGN_PARAGRAPH.CENTER, sp_before=4, sp_after=4)
        _edge_to_edge(p_ftr, inner_pad_cm=0)
        _hrule(p_ftr, pri, size=6, space=4, position="top")

    # ════════════════════════════════════════════════════════
    # BOLD  — dark shaded header, accent rule under it
    # ════════════════════════════════════════════════════════
    # ════════════════════════════════════════════════════════
    # BOLD  — full-width dark banner, accent rule under it, dark footer
    # ════════════════════════════════════════════════════════
    elif k == "bold":
        from docx.shared import Cm as _Cm
        # One-cell full-width table for the dark banner so text has proper
        # inner padding while the background spans the page edge-to-edge.
        tbl = hdr.add_table(rows=1, cols=1, width=_Cm(21))
        _no_borders(tbl)
        _full_width_table(tbl)
        tbl.columns[0].width = _Cm(21)
        cell = tbl.cell(0,0)
        _shade_cell(cell, sec_col)
        _set_cell_margins(cell, top=220, left=560, bottom=220, right=560)

        # Name paragraph
        p_name = cell.paragraphs[0]
        p_name.paragraph_format.space_after = Pt(2)
        r_logo = p_name.add_run()
        has_logo = _add_logo(r_logo, 0.48)
        if has_logo:
            p_name.add_run("  ")
        r_name = p_name.add_run(t.company_name or "")
        r_name.bold = True; r_name.font.name = fh; r_name.font.size = Pt(20)
        r_name.font.color.rgb = RGBColor(255,255,255)

        if t.tagline:
            p_tag = cell.add_paragraph()
            p_tag.paragraph_format.space_after = Pt(2)
            r_t = p_tag.add_run(t.tagline)
            r_t.italic = True; r_t.font.size = Pt(10)
            r_t.font.color.rgb = RGBColor(0xaa,0xaa,0xaa)

        if t.contact_info:
            p_c = cell.add_paragraph()
            p_c.paragraph_format.space_after = Pt(0)
            r_c = p_c.add_run(t.contact_info)
            r_c.font.size = Pt(9)
            r_c.font.color.rgb = RGBColor(0x99,0x99,0x99)

        # Accent rule below the banner — edge-to-edge
        p_acc = para(hdr, sp_before=0, sp_after=8)
        _edge_to_edge(p_acc, inner_pad_cm=0)
        _hrule(p_acc, pri, size=12, space=0)

        # Footer: dark band using a shaded table cell too
        ftbl = ftr.add_table(rows=1, cols=1, width=_Cm(21))
        _no_borders(ftbl)
        _full_width_table(ftbl)
        ftbl.columns[0].width = _Cm(21)
        fcell = ftbl.cell(0,0)
        _shade_cell(fcell, sec_col)
        _set_cell_margins(fcell, top=140, left=400, bottom=140, right=400)
        p_f = fcell.paragraphs[0]
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_f = p_f.add_run(f"{t.company_name or ''}  ·  {addr_flat}  ·  {t.contact_info or ''}")
        r_f.font.size = Pt(9)
        r_f.font.color.rgb = RGBColor(0x99,0x99,0x99)

    # ════════════════════════════════════════════════════════
    # MINIMAL  — name and contact inline, thin rule
    # ════════════════════════════════════════════════════════
    elif k == "minimal":
        from docx.shared import Cm as _Cm
        # 2-column table: name+tagline on left, address+contact on right,
        # both on the same baseline (matches the flex-row HTML preview).
        tbl = hdr.add_table(rows=1, cols=2, width=_Cm(16.0))
        _no_borders(tbl)
        # Use normal margins here (not full-width) so the content sits
        # inside the body column; the divider below goes edge-to-edge.
        tbl.columns[0].width = _Cm(7.0)
        tbl.columns[1].width = _Cm(9.0)
        cl = tbl.cell(0,0); cr = tbl.cell(0,1)
        _set_cell_margins(cl, top=40, left=0,   bottom=40, right=0)
        _set_cell_margins(cr, top=40, left=0,   bottom=40, right=0)
        _cell_vcenter(cl); _cell_vcenter(cr)

        # Left cell: logo + name on one line, tagline below (tagline on a
        # separate paragraph avoids it wrapping awkwardly after long names)
        p_name = cl.paragraphs[0]
        p_name.paragraph_format.space_after = Pt(1)
        r_logo = p_name.add_run()
        has_logo = _add_logo(r_logo, 0.42)
        if has_logo:
            p_name.add_run("  ")
        r_name = p_name.add_run(t.company_name or "")
        r_name.bold = True; r_name.font.name = fh; r_name.font.size = Pt(17)
        r_name.font.color.rgb = sec_rgb
        if t.tagline:
            p_tag = cl.add_paragraph()
            p_tag.paragraph_format.space_after = Pt(0)
            r_t = p_tag.add_run(t.tagline)
            r_t.font.size = Pt(10); r_t.italic = True
            r_t.font.color.rgb = RGBColor(0xbb,0xbb,0xbb)

        # Right cell: address + contact, right-aligned
        contact_line = addr_flat
        if t.contact_info:
            contact_line = f"{addr_flat}  ·  {t.contact_info}" if addr_flat else t.contact_info
        p_r = cr.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_r = p_r.add_run(contact_line)
        r_r.font.size = Pt(9)
        r_r.font.color.rgb = RGBColor(0xbb,0xbb,0xbb)

        # Divider rule spans the full page width
        p_div = para(hdr, sp_before=4, sp_after=0)
        _edge_to_edge(p_div, inner_pad_cm=0)
        _hrule(p_div, sec_col, size=4, space=2)

        para(ftr, f"{t.contact_info or ''}  |  {addr_flat}",
             size=9, color_rgb=RGBColor(0xbb,0xbb,0xbb),
             align=WD_ALIGN_PARAGRAPH.CENTER, sp_before=4, sp_after=4)

    # ════════════════════════════════════════════════════════
    # SPLIT  — two-column table: coloured left sidebar, white right
    # ════════════════════════════════════════════════════════
    elif k == "split":
        from docx.shared import Cm as _Cm
        tbl = hdr.add_table(rows=1, cols=2, width=_Cm(21))
        _no_borders(tbl)
        _full_width_table(tbl)
        # Left sidebar 5.5 cm, right content the rest
        tbl.columns[0].width = _Cm(5.5)
        tbl.columns[1].width = _Cm(15.5)

        cl = tbl.cell(0,0)
        cr = tbl.cell(0,1)
        _shade_cell(cl, pri)
        # Left cell: moderate top margin so the logo/name don't hug the
        # top edge of the sidebar (but the coloured bg still starts at top=0
        # because we don't change the table's top edge).
        _set_cell_margins(cl, top=280, left=280, bottom=280, right=280)
        _set_cell_margins(cr, top=0, left=280, bottom=280, right=400)

        # Left cell content — tight vertical spacing
        p_logo = cl.paragraphs[0]
        p_logo.paragraph_format.space_before = Pt(0)
        p_logo.paragraph_format.space_after = Pt(4)
        r_logo = p_logo.add_run()
        _add_logo(r_logo, 0.5)

        p_cl_name = cl.add_paragraph()
        p_cl_name.paragraph_format.space_after = Pt(2)
        r_cn = p_cl_name.add_run(t.company_name or "")
        r_cn.bold = True; r_cn.font.name = fh; r_cn.font.size = Pt(13)
        r_cn.font.color.rgb = RGBColor(255,255,255)

        if t.tagline:
            p_cl_tag = cl.add_paragraph()
            p_cl_tag.paragraph_format.space_after = Pt(0)
            r_ct = p_cl_tag.add_run(t.tagline)
            r_ct.font.size = Pt(9); r_ct.italic = True
            r_ct.font.color.rgb = RGBColor(0xdd,0xdd,0xdd)

        # Right cell — address + contact, right-aligned, vertically centred
        _cell_vcenter(cr)
        first_right_p = cr.paragraphs[0]   # first paragraph already exists
        first_right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        first_right_p.paragraph_format.space_after = Pt(0)
        if addr_lines:
            r = first_right_p.add_run(addr_lines[0])
            r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x88,0x88,0x88)
            for line in addr_lines[1:]:
                p_l = cr.add_paragraph()
                p_l.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p_l.paragraph_format.space_after = Pt(0)
                r = p_l.add_run(line)
                r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x88,0x88,0x88)
        if t.contact_info:
            p_rc = cr.add_paragraph()
            p_rc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_rc.paragraph_format.space_before = Pt(2)
            r = p_rc.add_run(t.contact_info)
            r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xaa,0xaa,0xaa)

        # Thin footer rule
        p_ftr = para(ftr, f"{t.contact_info or ''}  ·  {addr_flat}",
             size=9, color_rgb=RGBColor(0xaa,0xaa,0xaa),
             align=WD_ALIGN_PARAGRAPH.CENTER, sp_before=4, sp_after=4)
        _edge_to_edge(p_ftr, inner_pad_cm=0)
        _hrule(p_ftr, "DDDDDD", size=4, space=2, position="top")

    # ════════════════════════════════════════════════════════
    # ELEGANT  — centred logo, name, double rule, contact line
    # ════════════════════════════════════════════════════════
    elif k == "elegant":
        if _logo_path(t):
            p_logo = para(hdr, sp_before=4, sp_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
            r_logo = p_logo.add_run()
            _add_logo(r_logo, 0.55)

        p_name = para(hdr, t.company_name or "", bold=True, size=22, font=fh,
                      color_rgb=sec_rgb, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=3)
        if t.tagline:
            p_tag = para(hdr, t.tagline, italic=True, size=11, color_rgb=pri_rgb,
                         align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=6)

        # Double rule = thick rule then thin rule (both edge-to-edge)
        p_r1 = para(hdr, sp_before=2, sp_after=0)
        _edge_to_edge(p_r1, inner_pad_cm=0)
        _hrule(p_r1, pri, size=12, space=2)
        p_r2 = para(hdr, sp_before=0, sp_after=4)
        _edge_to_edge(p_r2, inner_pad_cm=0)
        _hrule(p_r2, pri, size=4, space=2)

        contact_line = addr_flat
        if t.contact_info:
            contact_line = f"{addr_flat}  ·  {t.contact_info}" if addr_flat else t.contact_info
        para(hdr, contact_line, size=9, color_rgb=RGBColor(0xaa,0xaa,0xaa),
             align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=4)

        # Double rule footer (edge-to-edge)
        p_f1 = para(ftr, sp_before=4, sp_after=0)
        _edge_to_edge(p_f1, inner_pad_cm=0)
        _hrule(p_f1, pri, size=12, space=2, position="top")
        para(ftr, t.contact_info or "", size=9, color_rgb=RGBColor(0xaa,0xaa,0xaa),
             align=WD_ALIGN_PARAGRAPH.CENTER, sp_before=2, sp_after=4)

    else:
        # Fallback to classic
        p_name = para(hdr, t.company_name or "", bold=True, size=18, font=fh, color_rgb=sec_rgb)
        p_div = para(hdr, sp_before=2, sp_after=0)
        _edge_to_edge(p_div, inner_pad_cm=0)
        _hrule(p_div, pri, size=6, space=2)
        para(ftr, t.contact_info or "", size=9, color_rgb=pri_rgb,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Body: render the user-authored HTML (from the builder) ──────────
    if getattr(t, 'body_html', None) and t.body_html.strip():
        _html_to_docx_paragraphs(t.body_html, doc, body_font=fb)
    else:
        # Default: one blank paragraph so the cursor has somewhere to land.
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════
# PDF  (reportlab — no LibreOffice)
# ═══════════════════════════════════════════════════════════════════════════
#
# IMPORTANT LAYOUT NOTE:
# ReportLab's coordinate system has (0,0) at the BOTTOM-LEFT of the page.
# The TOP of the page is y=H (≈841.89pt for A4). The HTML preview draws
# the letterhead header starting from the very top of the page with no
# outer margin, and the coloured accent shapes (top rule, banner, left
# panel) extend edge-to-edge. To match that, background shapes in this
# PDF export MUST draw from x=0..W and start at y=H (no subtraction of
# an outer margin). Only the text/logo content uses an inner horizontal
# pad (PAD_X) so it's not flush with the page edge.
#
def _build_pdf(t: LetterheadTemplate) -> bytes:
    _require('reportlab', 'reportlab', 'PDF export')
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.colors import HexColor, white, Color
    from reportlab.lib.utils import ImageReader

    buf = io.BytesIO()
    W, H = A4          # 595.27 × 841.89 pt

    # INNER horizontal pad for text/logo content (matches the HTML
    # preview's padding:14px 28px — at A4 ~595pt wide, 28px ≈ 28pt).
    PAD_X = 28.0

    pri = HexColor(t.color_primary  or "#1D9E75")
    sec = HexColor(t.color_secondary or "#2C2C2A")
    fh  = (t.font_header or "Georgia").strip("'")
    fb  = (t.font_body   or "Arial"  ).strip("'")
    k   = t.template_key or "classic"

    # Map CSS font names → ReportLab built-ins
    def rl_font(name, bold=False, italic=False):
        n = name.lower().strip("'")
        serif   = ("Times-Roman","Times-Bold","Times-Italic","Times-BoldItalic")
        mono    = ("Courier","Courier-Bold","Courier-Oblique","Courier-BoldOblique")
        sans    = ("Helvetica","Helvetica-Bold","Helvetica-Oblique","Helvetica-BoldOblique")
        if any(x in n for x in ("georgia","times","palatino","garamond")):
            fonts = serif
        elif any(x in n for x in ("courier","mono")):
            fonts = mono
        else:
            fonts = sans
        idx = (1 if bold else 0) + (2 if italic else 0)
        return fonts[min(idx,3)]

    font_h  = rl_font(fh, bold=True)
    font_h_ = rl_font(fh, bold=False)
    font_b  = rl_font(fb)
    font_bi = rl_font(fb, italic=True)

    c = rl_canvas.Canvas(buf, pagesize=A4)

    # Logo helper
    logo_img = None
    lp = _logo_path(t)
    if lp:
        try:
            logo_img = ImageReader(lp)
        except Exception:
            logo_img = None

    addr_lines = [l.strip() for l in (t.address or "").split("\n") if l.strip()]
    addr_flat  = "  ·  ".join(addr_lines)

    def draw_footer_band():
        """Coloured footer band edge-to-edge at the bottom of the page."""
        parts = []
        if t.contact_info: parts.append(t.contact_info)
        if addr_flat:      parts.append(addr_flat)
        text = "  ·  ".join(parts)
        c.setFillColor(pri)
        c.rect(0, 0, W, 28, stroke=0, fill=1)
        c.setFont(font_b, 8)
        c.setFillColor(white)
        c.drawCentredString(W/2, 10, text[:110])

    # ════════════════════════════════════════════════════════
    # CLASSIC — full-width top accent bar, then padded content
    # ════════════════════════════════════════════════════════
    if k == "classic":
        # Full-width top accent bar (5pt tall, flush to top of page)
        bar_h = 5
        c.setFillColor(pri)
        c.rect(0, H - bar_h, W, bar_h, stroke=0, fill=1)

        # Content starts just below the bar
        y = H - bar_h - 14     # ~14pt top padding inside the header

        # Logo + name (left side, with inner pad)
        lx = PAD_X
        if logo_img:
            lh_ = 42
            lw_ = 42
            c.drawImage(logo_img, lx, y - lh_, height=lh_, width=lw_,
                        preserveAspectRatio=True, mask="auto")
            lx += lw_ + 10

        c.setFont(font_h, 20); c.setFillColor(sec)
        c.drawString(lx, y - 20, t.company_name or "")

        if t.tagline:
            c.setFont(font_bi, 11); c.setFillColor(pri)
            c.drawString(lx, y - 34, t.tagline)

        # Address block on the right
        ay = y - 14
        c.setFont(font_b, 9); c.setFillColor(Color(0.53,0.53,0.53))
        for line in addr_lines:
            c.drawRightString(W - PAD_X, ay, line); ay -= 12
        if t.contact_info:
            c.setFillColor(Color(0.67,0.67,0.67))
            c.drawRightString(W - PAD_X, ay - 2, t.contact_info)

        # Thin divider under the header block
        div_y = y - 58
        c.setStrokeColor(Color(0.87,0.87,0.87)); c.setLineWidth(0.8)
        c.line(PAD_X, div_y, W - PAD_X, div_y)

        draw_footer_band()

    # ════════════════════════════════════════════════════════
    # MODERN — left accent stripe, padded content, thin rule, coloured footer text
    # ════════════════════════════════════════════════════════
    elif k == "modern":
        # Full-width region for the header (no top accent bar for modern)
        y = H - 14   # small top pad

        # Left accent stripe — short vertical bar from top edge
        stripe_w = 6
        stripe_h = 80
        c.setFillColor(pri)
        c.rect(0, H - stripe_h - 4, stripe_w, stripe_h, stroke=0, fill=1)

        lx = PAD_X
        if logo_img:
            c.drawImage(logo_img, lx, y - 42, height=38, width=38,
                        preserveAspectRatio=True, mask="auto")
            lx += 46

        c.setFont(font_h, 18); c.setFillColor(sec)
        c.drawString(lx, y - 18, t.company_name or "")
        if t.tagline:
            c.setFont(font_bi, 10); c.setFillColor(Color(0.67,0.67,0.67))
            c.drawString(lx, y - 32, t.tagline)

        ay = y - 14
        c.setFont(font_b, 9); c.setFillColor(Color(0.53,0.53,0.53))
        if t.contact_info:
            c.drawRightString(W - PAD_X, ay, t.contact_info); ay -= 12
        for line in addr_lines:
            c.drawRightString(W - PAD_X, ay, line); ay -= 12

        # Thin divider under the header
        div_y = y - 70
        c.setStrokeColor(Color(0.88,0.88,0.88)); c.setLineWidth(0.8)
        c.line(PAD_X, div_y, W - PAD_X, div_y)

        # Footer — accent line across, then coloured centred text
        parts = [t.contact_info or "", addr_flat]
        c.setStrokeColor(pri); c.setLineWidth(1.5)
        c.line(0, 32, W, 32)
        c.setFont(font_b, 8); c.setFillColor(pri)
        c.drawCentredString(W/2, 18, "  ·  ".join([p for p in parts if p])[:110])

    # ════════════════════════════════════════════════════════
    # BOLD — full-width dark banner at the very top, accent rule under it
    # ════════════════════════════════════════════════════════
    elif k == "bold":
        banner_h = 90 + (12 if t.tagline else 0) + (12 if t.contact_info else 0)
        # Full-width dark banner from top
        c.setFillColor(sec)
        c.rect(0, H - banner_h, W, banner_h, stroke=0, fill=1)

        # Accent rule right under the banner (full width)
        c.setFillColor(pri)
        c.rect(0, H - banner_h - 5, W, 5, stroke=0, fill=1)

        # Logo + name inside banner, padded from left
        lx = PAD_X
        ly_top = H - 16   # top of content inside banner
        if logo_img:
            c.drawImage(logo_img, lx, ly_top - 44, height=44, width=44,
                        preserveAspectRatio=True, mask="auto")
            lx += 54

        c.setFont(font_h, 20); c.setFillColor(white)
        c.drawString(lx, ly_top - 24, t.company_name or "")

        ty = ly_top - 42
        if t.tagline:
            c.setFont(font_bi, 10); c.setFillColor(Color(0.67,0.67,0.67))
            c.drawString(lx, ty, t.tagline)
            ty -= 14
        if t.contact_info:
            c.setFont(font_b, 9); c.setFillColor(Color(0.6,0.6,0.6))
            c.drawString(lx, ty, t.contact_info)

        # Dark footer band (matches banner)
        c.setFillColor(sec); c.rect(0, 0, W, 28, stroke=0, fill=1)
        c.setFont(font_b, 8); c.setFillColor(Color(0.6,0.6,0.6))
        footer_text = f"{t.company_name or ''}  ·  {addr_flat}  ·  {t.contact_info or ''}"
        c.drawCentredString(W/2, 10, footer_text[:110])

    # ════════════════════════════════════════════════════════
    # MINIMAL — name left, contact right, thin bottom rule (no top bar)
    # ════════════════════════════════════════════════════════
    elif k == "minimal":
        y = H - 14   # small top pad

        lx = PAD_X
        if logo_img:
            c.drawImage(logo_img, lx, y - 40, height=36, width=36,
                        preserveAspectRatio=True, mask="auto")
            lx += 44

        c.setFont(font_h, 17); c.setFillColor(sec)
        c.drawString(lx, y - 18, t.company_name or "")
        if t.tagline:
            c.setFont(rl_font(fh, italic=True), 10); c.setFillColor(Color(0.73,0.73,0.73))
            c.drawString(lx, y - 30, f"  {t.tagline}")

        contact_line = addr_flat
        if t.contact_info:
            contact_line = f"{addr_flat}  ·  {t.contact_info}" if addr_flat else t.contact_info
        c.setFont(font_b, 9); c.setFillColor(Color(0.73,0.73,0.73))
        c.drawRightString(W - PAD_X, y - 18, contact_line[:80])

        # Bottom rule in secondary colour (full width inside pad)
        r,g,b = _rgb(t.color_secondary or "#2C2C2A")
        c.setStrokeColor(Color(r/255,g/255,b/255)); c.setLineWidth(1.0)
        c.line(PAD_X, y - 50, W - PAD_X, y - 50)

        # Footer — thin rule + muted centred text
        c.setStrokeColor(Color(0.87,0.87,0.87)); c.setLineWidth(0.6)
        c.line(PAD_X, 28, W - PAD_X, 28)
        c.setFont(font_b, 8); c.setFillColor(Color(0.73,0.73,0.73))
        c.drawCentredString(W/2, 16, f"{t.contact_info or ''}   |   {addr_flat}"[:110])

    # ════════════════════════════════════════════════════════
    # SPLIT — coloured left panel edge-to-edge, white right side
    # ════════════════════════════════════════════════════════
    elif k == "split":
        panel_w = 180
        panel_h = 110

        # Left coloured panel: flush to top AND left edge (no outer margin)
        c.setFillColor(pri)
        c.rect(0, H - panel_h, panel_w, panel_h, stroke=0, fill=1)

        # Left panel content — logo + name (small inner pad)
        ly = H - 18
        if logo_img:
            c.drawImage(logo_img, 14, ly - 34, height=32, width=32,
                        preserveAspectRatio=True, mask="auto")
            ly -= 38

        c.setFont(font_h, 14); c.setFillColor(white)
        cname = (t.company_name or "")[:28]
        c.drawString(14, ly - 14, cname)
        if t.tagline:
            c.setFont(font_bi, 9); c.setFillColor(Color(0.87,0.87,0.87))
            c.drawString(14, ly - 28, t.tagline[:28])

        # Right address block (padded from right edge)
        ay = H - 22
        c.setFont(font_b, 9); c.setFillColor(Color(0.53,0.53,0.53))
        for line in addr_lines:
            c.drawRightString(W - PAD_X, ay, line); ay -= 12
        if t.contact_info:
            c.setFillColor(Color(0.67,0.67,0.67))
            c.drawRightString(W - PAD_X, ay - 2, t.contact_info)

        # Thin footer rule + muted text
        c.setStrokeColor(Color(0.87,0.87,0.87)); c.setLineWidth(0.6)
        c.line(PAD_X, 28, W - PAD_X, 28)
        c.setFont(font_b, 8); c.setFillColor(Color(0.67,0.67,0.67))
        c.drawCentredString(W/2, 16, f"{t.contact_info or ''}  ·  {addr_flat}"[:110])

    # ════════════════════════════════════════════════════════
    # ELEGANT — centred content with double rules (full-width rules)
    # ════════════════════════════════════════════════════════
    elif k == "elegant":
        y = H - 16   # small top pad

        if logo_img:
            c.drawImage(logo_img, W/2 - 24, y - 52, height=48, width=48,
                        preserveAspectRatio=True, mask="auto")
            y -= 56

        c.setFont(font_h, 22); c.setFillColor(sec)
        c.drawCentredString(W/2, y - 20, t.company_name or ""); y -= 22
        if t.tagline:
            c.setFont(font_bi, 11); c.setFillColor(pri)
            c.drawCentredString(W/2, y - 14, t.tagline); y -= 16

        # Double rule (full-width edge-to-edge)
        c.setStrokeColor(pri); c.setLineWidth(2.5)
        c.line(0, y - 20, W, y - 20)
        c.setLineWidth(0.8)
        c.line(0, y - 25, W, y - 25)

        contact_line = addr_flat
        if t.contact_info:
            contact_line = f"{addr_flat}  ·  {t.contact_info}" if addr_flat else t.contact_info
        c.setFont(font_b, 9); c.setFillColor(Color(0.67,0.67,0.67))
        c.drawCentredString(W/2, y - 40, contact_line[:100])

        # Double rule footer (full-width)
        c.setStrokeColor(pri); c.setLineWidth(2.5)
        c.line(0, 36, W, 36)
        c.setLineWidth(0.8)
        c.line(0, 31, W, 31)
        c.setFont(font_b, 8); c.setFillColor(Color(0.67,0.67,0.67))
        c.drawCentredString(W/2, 18, (t.contact_info or "")[:100])

    else:
        # Fallback — simple top bar + name + coloured footer
        bar_h = 4
        c.setFillColor(pri); c.rect(0, H - bar_h, W, bar_h, stroke=0, fill=1)
        y = H - bar_h - 24
        c.setFont(font_h, 18); c.setFillColor(sec)
        c.drawString(PAD_X, y, t.company_name or "")
        c.setStrokeColor(pri); c.setLineWidth(1)
        c.line(PAD_X, y - 10, W - PAD_X, y - 10)
        draw_footer_band()

    # ── Body: render the user-authored HTML (if any) ────────────────────
    # The body sits between the header zone (top ~130pt of the page) and
    # the footer band (bottom 28pt). We use a Frame + flowables so content
    # wraps naturally and honours the body font/colours.
    if getattr(t, 'body_html', None) and t.body_html.strip():
        try:
            from reportlab.platypus import Frame
            flowables = _html_to_pdf_flowables(t.body_html, pri, font_h, font_b)
            if flowables:
                # Header zones vary by style but ~135pt is a safe ceiling
                # that puts the body below every style's header.
                frame_top    = H - 155
                frame_bottom = 40        # clear the 28pt footer + breathing room
                frame_left   = PAD_X
                frame_right  = W - PAD_X
                frame = Frame(
                    frame_left, frame_bottom,
                    frame_right - frame_left,
                    frame_top   - frame_bottom,
                    showBoundary=0, leftPadding=0, rightPadding=0,
                    topPadding=4, bottomPadding=4,
                )
                frame.addFromList(flowables, c)
        except Exception as exc:
            logger.warning("PDF body render failed: %s", exc)

    c.save()
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════
# PNG  (Pillow — A4 portrait page at 150 DPI, 1240×1754 px)
# ═══════════════════════════════════════════════════════════════════════════
#
# The PNG export is a full A4 portrait page (not a landscape strip) so it
# can be used as a single-page image letterhead — drop it into any editor
# that supports image backgrounds. The header/footer layout mirrors the
# PDF export: background shapes go edge-to-edge; only text/logos respect
# an inner horizontal pad.

def _build_png(t: LetterheadTemplate) -> bytes:
    _require('PIL', 'Pillow', 'PNG export')
    from PIL import Image, ImageDraw, ImageFont

    # A4 portrait at 150 DPI — 210mm × 297mm
    W, H = 1240, 1754
    img  = Image.new("RGB", (W, H), (255,255,255))
    d    = ImageDraw.Draw(img)

    pri  = _rgb(t.color_primary  or "#1D9E75")
    sec_ = _rgb(t.color_secondary or "#2C2C2A")
    k    = t.template_key or "classic"
    MARG = 80   # inner horizontal pad for text/logo content

    # Font loading — try DejaVu first (common on Ubuntu), then Liberation
    def _font(size, bold=False, italic=False):
        candidates = []
        if bold and italic:
            candidates = ["/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf",
                          "/usr/share/fonts/truetype/liberation/LiberationSans-BoldItalic.ttf"]
        elif bold:
            candidates = ["/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                          "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"]
        elif italic:
            candidates = ["/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf",
                          "/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf"]
        else:
            candidates = ["/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                          "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"]
        for fp in candidates:
            if os.path.exists(fp):
                try: return ImageFont.truetype(fp, size)
                except: pass
        return ImageFont.load_default()

    def tw(text, font):
        try:
            bb = d.textbbox((0,0), text, font=font)
            return bb[2]-bb[0]
        except: return len(text)*7

    fn  = _font(38, bold=True)
    ft  = _font(20, italic=True)
    fb  = _font(18)
    fbs = _font(16)

    # Load logo
    logo_img = None
    lp = _logo_path(t)
    if lp:
        try:
            logo_img = Image.open(lp).convert("RGBA")
            lh = 70
            lw = int(lh * logo_img.width / logo_img.height)
            logo_img = logo_img.resize((lw, lh), Image.LANCZOS)
        except: logo_img = None

    addr_lines = [l.strip() for l in (t.address or "").split("\n") if l.strip()]
    addr_flat  = "  ·  ".join(addr_lines)

    def draw_footer_band(text, bg=None):
        bg = bg or pri
        d.rectangle([0, H-56, W, H], fill=bg)
        d.text(((W - tw(text, fbs))//2, H-38), text[:120], fill=(255,255,255), font=fbs)

    # ════════════════════════════════════════════════════════
    if k == "classic":
        d.rectangle([0, 0, W, 10], fill=pri)  # full-width top bar
        lx = MARG
        if logo_img:
            img.paste(logo_img, (lx, 40), logo_img)
            lx += logo_img.width + 14
        d.text((lx, 40), t.company_name or "", fill=sec_, font=fn)
        if t.tagline:
            d.text((lx, 86), t.tagline, fill=pri, font=ft)
        ay = 44
        for line in addr_lines:
            d.text((W - MARG - tw(line, fbs), ay), line, fill=(136,136,136), font=fbs); ay += 22
        if t.contact_info:
            d.text((W-MARG-tw(t.contact_info, fbs), ay+4), t.contact_info, fill=(170,170,170), font=fbs)
        d.rectangle([MARG, 132, W-MARG, 134], fill=(220,220,220))
        parts = [t.contact_info or "", addr_flat]
        draw_footer_band("  ·  ".join([p for p in parts if p]))

    elif k == "modern":
        d.rectangle([0, 0, 8, 130], fill=pri)  # flush-left accent stripe
        lx = MARG + 22
        if logo_img:
            img.paste(logo_img, (lx, 36), logo_img)
            lx += logo_img.width + 14
        d.text((lx, 40), t.company_name or "", fill=sec_, font=fn)
        if t.tagline:
            d.text((lx, 86), t.tagline, fill=(170,170,170), font=ft)
        ay = 44
        if t.contact_info:
            d.text((W-MARG-tw(t.contact_info, fbs), ay), t.contact_info, fill=(136,136,136), font=fbs); ay += 22
        for line in addr_lines:
            d.text((W-MARG-tw(line, fbs), ay), line, fill=(187,187,187), font=fbs); ay += 22
        d.rectangle([MARG+12, 134, W-MARG, 136], fill=(220,220,220))
        parts = [t.contact_info or "", addr_flat]
        d.rectangle([0, H-40, W, H], fill=(255,255,255))
        d.rectangle([0, H-40, W, H-38], fill=pri)
        text = "  ·  ".join([p for p in parts if p])
        d.text(((W - tw(text, fbs))//2, H-28), text[:120], fill=pri, font=fbs)

    elif k == "bold":
        banner_h = 160 + (30 if t.tagline else 0)
        d.rectangle([0, 0, W, banner_h], fill=sec_)
        lx = MARG
        if logo_img:
            img.paste(logo_img, (lx, 40), logo_img)
            lx += logo_img.width + 16
        d.text((lx, 44), t.company_name or "", fill=(255,255,255), font=fn)
        ty = 96
        if t.tagline:
            d.text((lx, ty), t.tagline, fill=(170,170,170), font=ft); ty += 34
        if t.contact_info:
            d.text((lx, ty), t.contact_info, fill=(153,153,153), font=fbs)
        d.rectangle([0, banner_h, W, banner_h+8], fill=pri)
        footer_text = f"{t.company_name or ''}  ·  {addr_flat}  ·  {t.contact_info or ''}"
        draw_footer_band(footer_text, bg=sec_)
        d.rectangle([0, H-56, W, H-53], fill=pri)  # accent line above footer

    elif k == "minimal":
        lx = MARG
        if logo_img:
            img.paste(logo_img, (lx, 34), logo_img)
            lx += logo_img.width + 14
        d.text((lx, 38), t.company_name or "", fill=sec_, font=fn)
        if t.tagline:
            tag_x = lx + tw(t.company_name or "", fn) + 20
            d.text((tag_x, 52), f"· {t.tagline}", fill=(187,187,187), font=ft)
        contact_line = addr_flat
        if t.contact_info: contact_line = f"{addr_flat}  ·  {t.contact_info}" if addr_flat else t.contact_info
        d.text((W-MARG-tw(contact_line, fbs), 48), contact_line[:80], fill=(187,187,187), font=fbs)
        d.rectangle([MARG, 118, W-MARG, 120], fill=sec_)
        d.rectangle([MARG, H-36, W-MARG, H-35], fill=(220,220,220))
        ftr = f"{t.contact_info or ''}   |   {addr_flat}"
        d.text(((W - tw(ftr, fbs))//2, H-28), ftr[:120], fill=(187,187,187), font=fbs)

    elif k == "split":
        panel_w = 300
        panel_h = 220   # fixed height — not tied to page height anymore
        d.rectangle([0, 0, panel_w, panel_h], fill=pri)
        ly = 26
        if logo_img:
            img.paste(logo_img, (20, ly), logo_img)
            ly += logo_img.height + 10
        name_lines = [t.company_name[i:i+18] for i in range(0, len(t.company_name or ""), 18)]
        for nl in name_lines[:2]:
            d.text((20, ly), nl, fill=(255,255,255), font=_font(26, bold=True)); ly += 34
        if t.tagline:
            d.text((20, ly+4), t.tagline[:22], fill=(221,221,221), font=_font(17, italic=True))
        ay = 40
        for line in addr_lines:
            d.text((W-MARG-tw(line, fbs), ay), line, fill=(136,136,136), font=fbs); ay += 22
        if t.contact_info:
            d.text((W-MARG-tw(t.contact_info, fbs), ay+4), t.contact_info, fill=(170,170,170), font=fbs)
        d.rectangle([0, H-56, W, H], fill=(240,240,240))
        ftr = f"{t.contact_info or ''}  ·  {addr_flat}"
        d.text(((W - tw(ftr, fbs))//2, H-38), ftr[:120], fill=(170,170,170), font=fbs)

    elif k == "elegant":
        y = 30
        if logo_img:
            img.paste(logo_img, ((W-logo_img.width)//2, y), logo_img)
            y += logo_img.height + 14
        name_text = t.company_name or ""
        d.text(((W - tw(name_text, fn))//2, y), name_text, fill=sec_, font=fn); y += 50
        if t.tagline:
            d.text(((W - tw(t.tagline, ft))//2, y), t.tagline, fill=pri, font=ft); y += 32
        # Double rule (full width)
        d.rectangle([0, y, W, y+4], fill=pri); y += 10
        d.rectangle([0, y, W, y+2], fill=pri); y += 12
        contact_line = addr_flat
        if t.contact_info: contact_line = f"{addr_flat}  ·  {t.contact_info}" if addr_flat else t.contact_info
        d.text(((W - tw(contact_line, fbs))//2, y), contact_line[:100], fill=(170,170,170), font=fbs)
        # Double rule footer (full width)
        d.rectangle([0, H-58, W, H-54], fill=pri)
        d.rectangle([0, H-50, W, H-48], fill=pri)
        ftr = t.contact_info or ""
        d.text(((W - tw(ftr, fbs))//2, H-38), ftr[:100], fill=(170,170,170), font=fbs)

    else:
        d.text((MARG, 40), t.company_name or "", fill=sec_, font=fn)
        d.rectangle([MARG, 90, W-MARG, 93], fill=pri)
        draw_footer_band(t.contact_info or "")

    out = io.BytesIO()
    img.save(out, format="PNG", dpi=(150, 150))
    return out.getvalue()


# ═══════════════════════════════════════════════════════════════════════════
# DJANGO VIEWS
# ═══════════════════════════════════════════════════════════════════════════

class LetterheadDocxExportView(LoginRequiredMixin, View):
    """GET /letterhead/api/templates/<uuid>/export/docx/"""
    def get(self, request, pk):
        template = get_object_or_404(LetterheadTemplate, pk=pk)
        try:
            data = _build_docx(template)
        except MissingDependencyError as exc:
            logger.error("DOCX dependency missing: %s", exc)
            return JsonResponse({"error": str(exc)}, status=501)
        except Exception as exc:
            logger.exception("DOCX export failed for %s", pk)
            return JsonResponse({"error": str(exc)}, status=500)
        safe = template.name.replace(" ","_").replace("/","-")
        r = HttpResponse(data,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        r["Content-Disposition"] = f'attachment; filename="letterhead_{safe}.docx"'
        r["Content-Length"] = len(data)
        return r


class LetterheadPdfExportView(LoginRequiredMixin, View):
    """GET /letterhead/api/templates/<uuid>/export/pdf/"""
    def get(self, request, pk):
        template = get_object_or_404(LetterheadTemplate, pk=pk)
        try:
            data = _build_pdf(template)
        except MissingDependencyError as exc:
            logger.error("PDF dependency missing: %s", exc)
            return JsonResponse({"error": str(exc)}, status=501)
        except Exception as exc:
            logger.exception("PDF export failed for %s", pk)
            return JsonResponse({"error": str(exc)}, status=500)
        safe = template.name.replace(" ","_").replace("/","-")
        r = HttpResponse(data, content_type="application/pdf")
        r["Content-Disposition"] = f'attachment; filename="letterhead_{safe}.pdf"'
        r["Content-Length"] = len(data)
        return r


class LetterheadPngExportView(LoginRequiredMixin, View):
    """GET /letterhead/api/templates/<uuid>/export/png/"""
    def get(self, request, pk):
        template = get_object_or_404(LetterheadTemplate, pk=pk)
        try:
            data = _build_png(template)
        except MissingDependencyError as exc:
            logger.error("PNG dependency missing: %s", exc)
            return JsonResponse({"error": str(exc)}, status=501)
        except Exception as exc:
            logger.exception("PNG export failed for %s", pk)
            return JsonResponse({"error": str(exc)}, status=500)
        safe = template.name.replace(" ","_").replace("/","-")
        r = HttpResponse(data, content_type="image/png")
        r["Content-Disposition"] = f'attachment; filename="letterhead_{safe}.png"'
        r["Content-Length"] = len(data)
        return r


# Keep this for backward-compat if old urls.py still references it
LetterheadDocxExportView = LetterheadDocxExportView